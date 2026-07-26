from io import BytesIO
import os

from django.conf import settings
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from dokumen.models import RiwayatJabatan, RiwayatPanggol, RiwayatPenempatan


BULAN = (
    '', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
    'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember',
)


def _safe(value, default='-'):
    text = str(value).strip() if value is not None else ''
    return text or default


def _set_cell_text(cell, value, *, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(_safe(value))
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    tc_pr.append(shading)


def _latest_employee_data(pegawai):
    jabatan = (
        RiwayatJabatan.objects.filter(pegawai=pegawai)
        .select_related(
            'nama_jabatan', 'jenjang_jabatan', 'unor', 'bidang',
            'sub_bidang', 'instalasi',
        )
        .order_by('-tmt_jabatan', '-created_at')
        .first()
    )
    panggol = (
        RiwayatPanggol.objects.filter(pegawai=pegawai)
        .select_related('panggol')
        .order_by('-tmt_gol', '-created_at')
        .first()
    )
    penempatan = (
        RiwayatPenempatan.objects.filter(pegawai=pegawai, status=True)
        .select_related(
            'penempatan_level1', 'penempatan_level2',
            'penempatan_level3', 'penempatan_level4',
        )
        .order_by('-updated_at')
        .first()
    )

    if jabatan:
        jabatan_lama = (
            jabatan.detail_nama_jabatan
            or jabatan.nama_jabatan
            or jabatan.jenjang_jabatan
        )
    else:
        jabatan_lama = '-'
    return {
        'jabatan_lama': _safe(jabatan_lama),
        'pangkat': _safe(panggol.panggol if panggol else None),
        'unit': _safe(penempatan.penempatan if penempatan else None),
    }


def _configure_normal_style(document):
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)


def _add_letter_header(document):
    table = document.add_table(rows=4, cols=4)
    table.autofit = False
    widths = (Cm(2.3), Cm(0.5), Cm(7.5), Cm(7.0))
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width

    labels = (
        ('Nomor', '${nomor_naskah}'),
        ('Sifat', 'Biasa'),
        ('Lampiran', '1 (satu) berkas'),
        ('Perihal', 'Usul Pengelolaan Jabatan Fungsional'),
    )
    for index, (label, value) in enumerate(labels):
        _set_cell_text(table.cell(index, 0), label, size=11)
        _set_cell_text(table.cell(index, 1), ':', size=11)
        _set_cell_text(table.cell(index, 2), value, size=11)
    _set_cell_text(
        table.cell(0, 3),
        'Lombok Tengah, ${tanggal_naskah}',
        size=11,
    )


def _add_letterhead(document):
    letterhead_path = os.path.join(
        settings.BASE_DIR, 'static', 'img', 'KOP RS MANDALIKA 2024.png'
    )
    if not os.path.exists(letterhead_path):
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(letterhead_path, width=Cm(16.5))


def _add_signature(document):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(10)
    table.columns[1].width = Cm(8)
    cell = table.cell(0, 1)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('Direktur\n\n${ttd_pengirim}\n\n').font.name = 'Arial'
    name = paragraph.add_run('dr. Oxy Tjahjo Wahjuni, Sp.EM., FICEP., FISQua')
    name.bold = True
    name.underline = True
    name.font.name = 'Arial'
    paragraph.add_run('\nPembina Tk. I (IV/b)\nNIP. 19710113 200112 2 001').font.name = 'Arial'


def generate_usulan_jabatan_docx(usulan_qs, periode):
    """Buat surat pengantar dan lampiran kolektif untuk satu periode."""
    usulan_list = list(usulan_qs)
    document = Document()
    _configure_normal_style(document)
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    _add_letterhead(document)
    _add_letter_header(document)
    document.add_paragraph()
    recipient = document.add_paragraph()
    recipient.add_run(
        'Yth. Gubernur Nusa Tenggara Barat\n'
        'Cq. Kepala Badan Kepegawaian Daerah\n'
        'di –\n'
        '    Mataram'
    )
    document.add_paragraph()
    document.add_paragraph('Dengan hormat,')
    body = document.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Cm(1.25)
    body.add_run(
        'Sehubungan dengan kebutuhan organisasi serta dalam rangka pengembangan '
        'karier Pegawai Negeri Sipil pada Rumah Sakit Mandalika Provinsi Nusa '
        'Tenggara Barat, bersama ini kami sampaikan usulan pengelolaan jabatan '
        'fungsional sesuai kertas kerja sebagaimana tercantum dalam lampiran.'
    )
    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    closing.paragraph_format.first_line_indent = Cm(1.25)
    closing.add_run(
        'Demikian surat pengantar ini kami sampaikan. Besar harapan kami agar '
        'usulan ini dapat diproses sesuai ketentuan yang berlaku. Atas perhatian '
        'dan kerja samanya, kami ucapkan terima kasih.'
    )
    document.add_paragraph()
    _add_signature(document)
    document.add_paragraph()
    document.add_paragraph(
        'Tembusan:\nKepala Dinas Kesehatan Provinsi Nusa Tenggara Barat'
    )

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    document.add_paragraph('Lampiran', style=None)
    info = document.add_table(rows=3, cols=3)
    for index, (label, value) in enumerate((
        ('Nomor', '${nomor_naskah}'),
        ('Tanggal', '${tanggal_naskah}'),
        ('Perihal', 'Usul Pengelolaan Jabatan Fungsional'),
    )):
        _set_cell_text(info.cell(index, 0), label, size=9)
        _set_cell_text(info.cell(index, 1), ':', size=9)
        _set_cell_text(info.cell(index, 2), value, size=9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        'DATA USULAN PENGELOLAAN JABATAN\n'
        f'PERIODE {BULAN[periode.month].upper()} {periode.year}'
    )
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    headers = (
        'No', 'Nama', 'NIP', 'Pangkat/Golru.', 'Kategori Pengelolaan',
        'Jabatan Lama', 'Jabatan Baru', 'Unit Kerja', 'SKP 2 Tahun',
        'Jabatan Lama Sesuai', 'PAK', 'Angka Kredit', 'Sertifikat Uji Komp.',
        'Formasi Tersedia',
    )
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        _set_cell_text(
            table.cell(0, index), header, bold=True, size=7,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_shading(table.cell(0, index), 'D9EAF7')

    for number, usulan in enumerate(usulan_list, 1):
        profil = getattr(usulan.pegawai, 'profil_user', None)
        employee_data = _latest_employee_data(usulan.pegawai)
        row = table.add_row()
        values = (
            number,
            usulan.pegawai.full_name_2,
            profil.nip if profil else '-',
            employee_data['pangkat'],
            usulan.get_kategori_pengelolaan_display(),
            employee_data['jabatan_lama'],
            usulan.jabatan_diusulkan,
            employee_data['unit'],
            'Ya' if usulan.kinerja_dua_thn.count() == 2 else 'Tidak',
            'Sesuai',
            '√' if usulan.pak_id else '-',
            usulan.pak.ak if usulan.pak_id else '-',
            '√' if usulan.kompetensi_id else '-',
            '√' if usulan.formasi_tersedia else '-',
        )
        for index, value in enumerate(values):
            _set_cell_text(
                row.cells[index], value, size=7,
                align=WD_ALIGN_PARAGRAPH.CENTER if index not in (1, 5, 6, 7) else WD_ALIGN_PARAGRAPH.LEFT,
            )

    document.add_paragraph()
    _add_signature(document)
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
