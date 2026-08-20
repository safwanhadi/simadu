from django.shortcuts import render, redirect
from django.views import View
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.exceptions import PermissionDenied
from django.urls import reverse
from django.http import FileResponse, Http404, HttpResponse
from django.contrib import messages
from django.utils import timezone
from django.utils.text import slugify
from strukturorg.services import get_active_leader
from datetime import datetime, date, timedelta
from dateutil.relativedelta import relativedelta
from django.shortcuts import get_object_or_404
import qrcode
from io import BytesIO
import locale
import json
import inspect

#python docx --> library for generate docx
from num2words import num2words
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
from docx import Document as CreateDocument
from docx.document import Document
from docx.shared import Inches, Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from reportlab.graphics.barcode import qr


#reportlab --> libary for generate pdf
from reportlab.graphics.shapes import Drawing
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib.units import mm, cm
from reportlab.lib.pagesizes import A4, portrait, inch, legal
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
)
from dokumen.utils.jabatan import get_jabatan_unit

from layanan.models import (
    LayananGajiBerkala, LayananCuti, LayananUsulanInovasi, VerifikasiCuti, PelimpahanTugas
)
from dokumen.models import (
    RiwayatCuti,
    RiwayatDiklat,
    RiwayatJabatan,
    RiwayatPenempatan,
    RiwayatPengangkatan,
)
from .models import TextSPTDiklat
from .forms import TextSPTDiklatForm
from layanan.services import CheckCuti
from layanan.utils import get_nip
from layanan.access.cuti import (
    build_approval_chain,
    can_view_delegation,
    can_view_leave,
)

locale.setlocale(locale.LC_ALL, 'id_ID.utf-8')


class CutiFileDownloadView(LoginRequiredMixin, View):
    """Mengirim berkas cuti setelah pemeriksaan hak akses."""

    allowed_fields = {'file_pengajuan', 'file_pendukung', 'file'}

    def get(self, request, pk, field_name):
        if field_name not in self.allowed_fields:
            raise Http404
        riwayat = get_object_or_404(
            RiwayatCuti.objects.select_related('usulan', 'pegawai'),
            pk=pk,
        )
        if riwayat.usulan is None or not can_view_leave(request.user, riwayat.usulan):
            raise PermissionDenied("Anda tidak berhak mengunduh berkas cuti ini.")
        field_file = getattr(riwayat, field_name)
        if not field_file:
            raise Http404
        return FileResponse(
            field_file.open('rb'),
            as_attachment=True,
            filename=field_file.name.rsplit('/', 1)[-1],
        )

def clear_document(doc):
    for element in doc.element.body:
        doc.element.body.remove(element)
    
def get_string_date_from_datetime(tanggal):
    tanggal_sekarang = date.today()
    try:
        get_tanggal = datetime.strftime(tanggal, "%d %B %Y")
        return get_tanggal
    except Exception:
        return datetime.strftime(tanggal_sekarang, "%d %B %Y")


#generate pdf
def capaiananak(request, kat, bulan, tahun):
    response = HttpResponse(content_type='application/pdf')
    today = datetime.today()
    d = today.strftime('%Y-%m-%d')
    response['Content-Disposition'] = f'inline: filename="{d}.pdf"'
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    p.setPageSize(portrait(A4))
    p.setTitle(f'Capaian Program Anak - Diunduh - {d}')
    p.drawCentredString(400, 510, 'LAPORAN CAPAIAN PROGRAM KESEHATAN ANAK')
    p.line(50, 500, 760, 500)


alignment_dict = {'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                  'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                  'centre': WD_PARAGRAPH_ALIGNMENT.CENTER,
                  'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                  'left': WD_PARAGRAPH_ALIGNMENT.LEFT}

document:Document = CreateDocument()
def add_content(doc=CreateDocument(), content="", content2="", style=None, space_before=0, space_after=0, tab=0, font_name="Times New Roman", font_size=12,
                set_bold=False, set_italic=False, set_underline=False, align="justify", line_spacing=0, keep_together=False, keep_with_next=False, 
                page_break_before=False, widow_control=False, left_indent=0):
    paragraph = doc.add_paragraph(content, style=style)
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = set_bold
    font.italic = set_italic
    font.underline = set_underline
    paragraph.paragraph_format.alignment = alignment_dict.get(align.lower())
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    paragraph.paragraph_format.keep_together = keep_together
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.paragraph_format.widow_control = widow_control
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(tab))
    paragraph.paragraph_format.left_indent = Mm(left_indent)
    data = paragraph.add_run(content2)
    data.bold = set_bold
    data.underline = set_underline
    data.italic = set_italic

def add_table_content(table=document.add_table(0,0), row_index=0, col_index=0, content="", content2="", style=None, space_before=0, space_after=0, tab=0, font_name="Times New Roman", font_size=12,
                set_bold=False, set_italic=False, set_underline=False, align="left", line_spacing=0, keep_together=False, keep_with_next=False, 
                page_break_before=False, widow_control=False, left_indent=0):
    font = table.style.font
    font.name = font_name
    font.size = Pt(font_size)
    paragraph = table.cell(row_idx=row_index, col_idx=col_index).add_paragraph(content, style=style)
    paragraph.paragraph_format.alignment = alignment_dict.get(align.lower())
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    paragraph.paragraph_format.keep_together = keep_together
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.paragraph_format.widow_control = widow_control
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(tab))
    paragraph.paragraph_format.left_indent = Mm(left_indent)
    data = paragraph.add_run(content2)
    data.bold = set_bold
    data.underline = set_underline
    data.italic = set_italic

# table=doc.add_table(0,0)
def add_table_content2(table=document.add_table(0,0), row_index=0, col_index=0, content="", style='Table Grid', space_before=0, space_after=0, tab=0, font_name="Times New Roman", font_size=12,
                set_bold=False, set_italic=False, set_underline=False, align="left", line_spacing=0, keep_together=False, keep_with_next=False, 
                page_break_before=False, widow_control=False, left_indent=0):
    table.style=style
    font = table.style.font
    font.name = font_name
    font.size = Pt(font_size)
    cell = table.cell(row_idx=row_index, col_idx=col_index)
    paragraph = cell.paragraphs[0]
    data = paragraph.add_run(content)
    data.bold = set_bold
    data.underline = set_underline
    data.italic = set_italic
    data.font.size = Pt(font_size)
    data.font.name = font_name
    # paragraph.add_run().add_picture()
    paragraph.paragraph_format.alignment = alignment_dict.get(align.lower())
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    paragraph.paragraph_format.keep_together = keep_together
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.paragraph_format.widow_control = widow_control
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(tab))
    paragraph.paragraph_format.left_indent = Mm(left_indent)


def convert_num_2_word(angka):
        number_dict = {i: num2words(i, lang='id') for i in range(0, 360)}
        return number_dict.get(angka, "Tidak dikenali")

    
class LayananGajiBerkalaDocxView(View):
    def set_col_widths(self, table):
        widths = (Inches(1), Inches(0.2), Inches(2.5), Inches(1), Inches(3))
        # height = (Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                # row.cells[idx].height = height

    def set_paragraph_space(self, row):
        for cell in row:
            for content in cell.paragraphs:
                content.paragraph_format.space_after=Pt(0.0)
                content.paragraph_format.space_before=Pt(0.0)

    def add_table_content(self, row, contents):
        # for cell in row:
        for idx, content in enumerate(contents):
            row[idx].text = content

    def get_object(self, id):
        try:
            data = LayananGajiBerkala.objects.get(id=id)
            return data
        except Exception:
            return None
        
    def get_unor_pimpinan(self, pegawai, layanan_cuti=None):
        try:
            penempatan = RiwayatPenempatan.objects.filter(pegawai=pegawai, status=True).order_by('-id').first()
            if penempatan:
                data = penempatan.unor_pimpinan
                snapshot = getattr(layanan_cuti, 'verifikasicuti', None)
                if snapshot:
                    signer = next((user for user in (
                        snapshot.verifikator3,
                        snapshot.verifikator2,
                        snapshot.verifikator1,
                    ) if user is not None), None)
                    if signer:
                        data['nama_pimpinan'] = getattr(signer, 'full_name_2', 'N/A')
                        profil = getattr(signer, 'profil_user', None)
                        data['nip'] = getattr(profil, 'nip', 'N/A') or 'N/A'
                        panggol = signer.riwayatpanggol_set.order_by('-id').first()
                        data['panggol'] = getattr(panggol, 'panggol', 'N/A') or 'N/A'
                        struktur = penempatan.unor
                        masa_jabatan = struktur.riwayat_pejabat.filter(
                            pejabat=signer,
                        ).order_by('-tanggal_mulai', '-id').first() if struktur else None
                        if masa_jabatan and masa_jabatan.nama_jabatan:
                            data['pimpinan'] = masa_jabatan.nama_jabatan
                return data
            return ""
        except Exception:
            return ""
        

    def get(self, request, **kwargs):
        #Pengaturan ukuran kertas
        doc:Document=CreateDocument()
        id = kwargs.get('layanan_id')
        layanan_berkala = self.get_object(id=id)
        page_size = doc.sections[0]
        page_size.page_width = Mm(215.9)
        page_size.page_height = Mm(330.0)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.bottom_margin = Inches(1.0)
        #penambahan gambar kop
        doc.add_picture('static/img/KOP RS MANDALIKA 2024.png', width=Inches(7.0))

        table1 = doc.add_table(rows=5, cols=5)
        table1.autofit = False
        self.set_col_widths(table1)
        
        row1 = table1.rows[0].cells
        contents = ('','','','','Kepada:')
        self.add_table_content(row1, contents)
        self.set_paragraph_space(row1)

        row2 = table1.rows[1].cells
        contents = ('Nomor', ':', '${nomor_naskah}', '', 'Yth. Kepala Badan Pengelolaan')
        self.add_table_content(row2, contents)
        self.set_paragraph_space(row2)

        row3 = table1.rows[2].cells
        contents = ('Sifat', ':', '${sifat}', '', 'Keuangan dan Aset Daerah')
        self.add_table_content(row3, contents)
        self.set_paragraph_space(row3)

        row4 = table1.rows[3].cells
        contents = ('Lampiran', ':', '-', '', 'Provinsi NTB')
        self.add_table_content(row4, contents)
        self.set_paragraph_space(row4)

        row5 = table1.rows[4].cells
        contents = ('Perihal', ':', f'Kenaikan Gaji Berkala a/n {layanan_berkala.pegawai.full_name_2 if layanan_berkala is not None else ""}, NIP. {layanan_berkala.pegawai.profil_user.nip  if layanan_berkala is not None else ""}', '', 'di Mataram')
        self.add_table_content(row5, contents)
        self.set_paragraph_space(row5)

        paragraph1 = doc.add_paragraph('Dengan ini diberitahukan bahwa berhubung dengan telah dipenuhinya masa kerja dan syarat-syarat lainnya kepada:')
        paragraph1.paragraph_format.space_before = Pt(16)
        paragraph1.paragraph_format.space_after = Pt(6)

        add_content(doc=doc, content=f'Nama\t: {layanan_berkala.pegawai.full_name_2 if layanan_berkala is not None and hasattr(layanan_berkala.pegawai, "full_name") else ""}', style='List Number', tab=2.5, left_indent=7)
        add_content(doc=doc, content=f'NIP\t: {layanan_berkala.pegawai.profil_user.nip if layanan_berkala is not None and hasattr(layanan_berkala.pegawai, "profil_user") else ""}', style="List Number", tab=2.5, left_indent=7)
        add_content(doc=doc, content=f'Pangkat\t: {layanan_berkala.riwayat.pangkat.panggol if layanan_berkala is not None and hasattr(layanan_berkala.riwayat, "pangkat") and hasattr(layanan_berkala.riwayat.pangkat, "panggol") else ""}', style="List Number", tab=2.5, left_indent=7)
        add_content(doc=doc, content=f'Tempat Bekerja\t: {layanan_berkala.riwayat.tempat_kerja.unor if layanan_berkala is not None and hasattr(layanan_berkala.riwayat, "tempat_kerja") and hasattr(layanan_berkala.riwayat.tempat_kerja, "unor") else ""}', style="List Number", tab=2.5, left_indent=7)
        add_content(doc=doc, content=f'Gaji Pokok\t: {locale.currency(layanan_berkala.riwayat.gaji_pkk, grouping=True) if layanan_berkala is not None and hasattr(layanan_berkala.riwayat, "gaji_pkk") else ""} {layanan_berkala.riwayat.pertek}', style="List Number", tab=2.5, left_indent=7)
        add_content(doc=doc, content='Atas dasar keputusan gaji/pangkat yang ditetapkan : ', style="List Number", space_after=6, left_indent=7)
        add_content(doc=doc, content='\ta.   Oleh', tab=0.3, content2=f' Pejabat\t\t\t: {layanan_berkala.riwayat.tempat_kerja.pimpinan if layanan_berkala is not None and hasattr(layanan_berkala.riwayat, "tempat_kerja") and hasattr(layanan_berkala.riwayat.tempat_kerja, "pimpinan") else ""}', left_indent=7)
        add_content(doc=doc, content='\tb.   Nomor dan', tab=0.3, content2=f' Tanggal\t\t: {layanan_berkala.riwayat.no_srt_gaji if layanan_berkala is not None and hasattr(layanan_berkala.riwayat, "no_srt_gaji") else ""}', left_indent=7)
        add_content(doc=doc, content='\tc.   Terhitung mulai', tab=0.3, content2=f' tanggal\t: {layanan_berkala.riwayat.tmt_gaji.strftime("%d %B %Y") if layanan_berkala is not None and hasattr(layanan_berkala.riwayat, "tmt_gaji") else ""}', left_indent=7)
        add_content(doc=doc, content=f'\td.   Masa kerja golongan pada tanggal tersebut : {layanan_berkala.riwayat.masa_kerja_tahun if layanan_berkala is not None and hasattr(layanan_berkala.riwayat, "masa_kerja_tahun") else ""} Tahun {layanan_berkala.riwayat.masa_kerja_bulan if layanan_berkala is not None and hasattr(layanan_berkala.riwayat, "masa_kerja_bulan") else ""} Bulan', tab=0.3, left_indent=7)
        add_content(doc=doc, content2='Diberikan kenaikan gaji berkala hingga memperoleh :', set_bold=True, set_underline=True, space_after=6, space_before=10)
        add_content(doc=doc, content=f'Gaji Pokok Baru\t: {locale.currency(layanan_berkala.berkala.gaji_pkk, grouping=True) if layanan_berkala is not None and hasattr(layanan_berkala.berkala, "gaji_pkk") else ""}', style='List Number', tab=2.5, left_indent=7)
        add_content(doc=doc, content=f'Berdasarkan masa kerja\t: {layanan_berkala.berkala.masa_kerja_tahun if layanan_berkala is not None and hasattr(layanan_berkala.berkala, "masa_kerja_tahun") else ""} Tahun {layanan_berkala.berkala.masa_kerja_bulan if layanan_berkala is not None and hasattr(layanan_berkala.berkala, "masa_kerja_bulan") else ""} Bulan', style='List Number', tab=2.5, left_indent=7)
        add_content(doc=doc, content=f'Dalam golongan ruang\t: {layanan_berkala.berkala.pangkat.panggol if layanan_berkala is not None and hasattr(layanan_berkala.berkala, "pangkat") else ""}', style='List Number', tab=2.5, left_indent=7)
        add_content(doc=doc, content=f'Mulai tanggal\t: {layanan_berkala.berkala.tmt_gaji.strftime("%d %B %Y") if layanan_berkala is not None and hasattr(layanan_berkala.berkala, "tmt_gaji") else ""}', style='List Number', tab=2.5, left_indent=7)
        add_content(doc=doc, content=f'Diharapkan agar sesuai {layanan_berkala.berkala.pertek if layanan_berkala is not None and hasattr(layanan_berkala.berkala, "pertek") else "" } kepada pegawai tersebut dapat dibayarkan penghasilannya berdasarkan gaji pokok yang baru.', space_after=10, space_before=6)
        add_content(doc=doc, content='Lombok Tengah, ${tanggal_naskah}', align='left', left_indent=80)
        add_content(doc=doc, content=f'{layanan_berkala.riwayat.tempat_kerja.pimpinan if layanan_berkala is not None and hasattr(layanan_berkala.riwayat, "tempat_kerja") and hasattr(layanan_berkala.riwayat.tempat_kerja, "pimpinan") else ""}', left_indent=80, align='left', space_after=30)
        add_content(doc=doc, content='${ttd_pengirim}', left_indent=80, align='left', space_after=30)
        add_content(doc=doc, content=f'{self.get_unor_pimpinan(layanan_berkala.pegawai)["nama_pimpinan"] if self.get_unor_pimpinan(layanan_berkala.pegawai) else ""}', left_indent=80, align='left')
        add_content(doc=doc, content=f'{self.get_unor_pimpinan(layanan_berkala.pegawai)["panggol"] if self.get_unor_pimpinan(layanan_berkala.pegawai) else ""}', left_indent=80, align='left')
        add_content(doc=doc, content=f'{self.get_unor_pimpinan(layanan_berkala.pegawai)["nip"] if self.get_unor_pimpinan(layanan_berkala.pegawai) else ""}', left_indent=80, align='left')
        add_content(doc=doc, content='Tembusan :', space_before=30, )
        add_content(doc=doc, content='1.Ketua Badan Pemeriksaan Keuangan di Jakarta')
        add_content(doc=doc, content='2.Mendagri (Kepala Biro Kepeg./Kepala Bag. Mutasi Peg. Daerah) di Jakarta')
        add_content(doc=doc, content='3.Kepala Badan Kepegawaian Negara di Jakarta')
        add_content(doc=doc, content='4.Kepala Kantor Regional X BKN di Denpasar')
        add_content(doc=doc, content='5.Kepala Badan Kepegawaian Daerah Provinsi NTB di Mataram')
        add_content(doc=doc, content='6.Inspektur Inspektorat Provinsi NTB di Mataram')
        add_content(doc=doc, content='7.Ka. Sub Bag Keuangan/Pembuat Daftar Gaji Dikes Provivinsi NTB di Mataram')
        add_content(doc=doc, content='8.Yang Bersangkutan Untuk Maklum')

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=berkala-{layanan_berkala.pegawai.first_name}.docx'
        doc.save(response)

        return response


class LayananUsulanCutiDocxView(LoginRequiredMixin, View):
    def set_col_widths(self, table):
        widths = (Inches(3.5), Inches(4.5))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

    def get_object(self, id):
        try:
            data = LayananCuti.objects.get(id=id)
            return data
        except Exception:
            return None
    
    def get_data_cuti(self, data_input=None, data_output=None, attr=""):
        if data_input is not None and hasattr(data_input.cuti_usulan, attr):
            return data_output
        return ""

    def get_verification_object(self, id):
        try:
            cuti = LayananCuti.objects.get(id=id)
            data = VerifikasiCuti.objects.get(layanan_cuti=cuti)
            return data
        except LayananCuti.DoesNotExist:
            return None

    def get(self, request, **kwargs):
        #Pengaturan ukuran kertas
        id = kwargs.get('layanan_id')
        doc:Document=CreateDocument()
        verifikasi_cuti = self.get_verification_object(id)
        layanan_cuti = self.get_object(id=id)
        if layanan_cuti is None or not can_view_leave(request.user, layanan_cuti):
            raise PermissionDenied("Anda tidak berhak mengunduh dokumen cuti ini.")
        verifikator1 = None
        verifikator2 = None
        verifikator3 = None
        sdm_cuti = None
        if layanan_cuti is not None:
            sdm_cuti = f"""
            pegawai: {layanan_cuti.pegawai.full_name_2}
            nip: {layanan_cuti.pegawai.profil_user.nip if layanan_cuti is not None and hasattr(layanan_cuti.pegawai, 'profil_user') else None}
            validasi_file: http://{request.get_host()}{reverse("file_urls:usulan_cuti_docx", kwargs={"layanan_id":id})}
            """
        if verifikasi_cuti is not None:
            verifikator1 = f"""
            id : {verifikasi_cuti.id} 
            verifikator : {verifikasi_cuti.verifikator1.full_name_2 if verifikasi_cuti.verifikator1 is not None else ""}
            nip : {verifikasi_cuti.verifikator1.profil_user.nip if verifikasi_cuti.verifikator1 is not None and hasattr(verifikasi_cuti.verifikator1, 'profil_user') else ""}
            validasi_file : http://{request.get_host()}{reverse("file_urls:usulan_cuti_docx", kwargs={"layanan_id":id})}
            """
            verifikator2 = f"""
            id : {verifikasi_cuti.id} 
            verifikator : {verifikasi_cuti.verifikator2.full_name_2 if verifikasi_cuti.verifikator2 is not None else ""}
            nip : {verifikasi_cuti.verifikator2.profil_user.nip if verifikasi_cuti.verifikator2 is not None and hasattr(verifikasi_cuti.verifikator2, 'profil_user') else ""}
            validasi_file : https://{request.get_host()}{reverse("file_urls:usulan_cuti_docx", kwargs={"layanan_id":id})}
            """
            verifikator3 = f"""
            id : {verifikasi_cuti.id} 
            verifikator : {verifikasi_cuti.verifikator3.full_name_2 if verifikasi_cuti.verifikator3 is not None else ""}
            nip : {verifikasi_cuti.verifikator3.profil_user.nip if verifikasi_cuti.verifikator3 is not None and hasattr(verifikasi_cuti.verifikator3, 'profil_user') else ""}
            """
        jabatan = layanan_cuti.pegawai.riwayatjabatan_set.last()
        data_instansi = layanan_cuti.pegawai.riwayat_penempatan.last()
        lama_cuti = layanan_cuti.cuti_usulan.lama_cuti if layanan_cuti is not None and hasattr(layanan_cuti, "cuti_usulan") else 0
        page_size = doc.sections[0]
        page_size.page_width = Mm(215.9)
        page_size.page_height = Mm(330.0)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.bottom_margin = Inches(1.0)
        #penambahan gambar kop
        doc.paragraphs.clear()
        doc.add_picture('static/img/KOP RS MANDALIKA 2024.png', width=Inches(7.0))
        add_content(doc=doc, content2='LAMPIRAN IV SURAT EDARAN KEPALA BADAN ADMINISTRASI KEPEGAWAIAN NEGARA', align='left', left_indent=75.0, set_bold=True)
        add_content(doc=doc, content2='NOMOR \t: 893/1450/BKD/2018', left_indent=90.0, tab=1, set_bold=True)
        add_content(doc=doc, content2='TANGGAL \t: 21 Mei 2018', left_indent=90.0, tab=1, set_bold=True, space_after=10)
        add_content(doc=doc, content=f'Lombok Tengah, {verifikasi_cuti.tanggal.strftime("%d %B %Y") if verifikasi_cuti.tanggal is not None else ""}', align='right', space_after=6)

        table1 = doc.add_table(rows=1, cols=2)
        table1.autofit = False
        self.set_col_widths(table1)
        add_table_content2(table=table1, row_index=0, col_index=0, content=f'PERMINTAAN {(self.get_data_cuti(layanan_cuti, layanan_cuti.cuti_usulan.jenis_cuti, "jenis_cuti")).upper()}', font_size=13, set_bold=True)
        add_table_content2(table=table1, row_index=0, col_index=1, style=None, content='\tK e p a d a \nYth.\tDirektur RS Mandalika  \n\tProvinsi Nusa Tenggara Barat \n\tdi Lombok Tengah')
        # generate qrcode
        pegawai_qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=7, border=4,)
        pegawai_qr.add_data(sdm_cuti)
        pegawai_qr.make(fit=True)
        img = pegawai_qr.make_image(fill='black', back_color='white')
        # Save QR code to a BytesIO object
        pegawai_buffer = BytesIO()
        img.save(pegawai_buffer, format='PNG')
        pegawai_buffer.seek(0)

        kasi_qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=7, border=4,)
        kasi_qr.add_data(verifikator1)
        kasi_qr.make(fit=True)
        img = kasi_qr.make_image(fill='black', back_color='white')
        kasi_buffer = BytesIO()
        img.save(kasi_buffer, format='PNG')
        kasi_buffer.seek(0)

        kabid_qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=7, border=4,)
        kabid_qr.add_data(verifikator2)
        kabid_qr.make(fit=True)
        img = kabid_qr.make_image(fill='black', back_color='white')
        kabid_buffer = BytesIO()
        img.save(kabid_buffer, format='PNG')
        kabid_buffer.seek(0)

        dir_qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=7, border=4,)
        dir_qr.add_data(verifikator3)
        dir_qr.make(fit=True)
        img = dir_qr.make_image(fill='black', back_color='white')
        dir_buffer = BytesIO()
        img.save(dir_buffer, format='PNG')
        dir_buffer.seek(0)
        
        add_content(doc=doc, content='Yang bertanda tangan dibawah ini :', space_before=10, space_after=6)
        add_content(doc=doc, content=f'\tNama \t\t\t:  {layanan_cuti.pegawai.full_name_2 if layanan_cuti is not None and hasattr(layanan_cuti.pegawai, "full_name_2") else ""}', align="left")
        add_content(doc=doc, content=f'\tNIP/NRK \t\t: {layanan_cuti.pegawai.profil_user.nip if layanan_cuti is not None and hasattr(layanan_cuti.pegawai, "profil_user") and hasattr(layanan_cuti.pegawai.profil_user, "nip") else ""}')
        add_content(doc=doc, content=f'\tJabatan \t\t: {jabatan.nama_jabatan if jabatan is not None and hasattr(jabatan, "nama_jabatan") else ""}')
        add_content(doc=doc, content=f'\tSatuan Organisasi \t: {data_instansi.unor if data_instansi is not None else "" }')
        add_content(doc=doc, content=f'Dengan ini mengajukan permintaan {self.get_data_cuti(layanan_cuti, layanan_cuti.cuti_usulan.jenis_cuti, "jenis_cuti")} karena {self.get_data_cuti(layanan_cuti, layanan_cuti.cuti_usulan.alasan_cuti, "alasan_cuti")} selama {convert_num_2_word(lama_cuti)} ({lama_cuti}) hari, terhitung mulai tanggal {self.get_data_cuti(layanan_cuti, layanan_cuti.cuti_usulan.tgl_mulai_cuti.strftime("%d %B"), "tgl_mulai_cuti")} s.d {self.get_data_cuti(layanan_cuti, layanan_cuti.cuti_usulan.tgl_akhir_cuti.strftime("%d %B %Y"), "tgl_akhir_cuti")} dan selama menjalankan cuti alamat saya adalah {self.get_data_cuti(layanan_cuti, layanan_cuti.cuti_usulan.domisili_saat_cuti, "domisili_saat_cuti")}', space_after=6, space_before=6)
        add_content(doc=doc, content='Demikian permintaan ini saya buat untuk dapat dipertimbangkan sebagaimana mestinya.', space_after=16)
        add_content(doc=doc, content=f'Hormat Saya', left_indent=80, align='center')
        paragraf = doc.add_paragraph()
        paragraf.add_run().add_picture(pegawai_buffer, Inches(1))
        paragraf.paragraph_format.left_indent=Mm(80.0)
        paragraf.paragraph_format.alignment = alignment_dict.get('center')
        add_content(doc=doc, content2=f'{layanan_cuti.pegawai.full_name_2 if layanan_cuti is not None and hasattr(layanan_cuti.pegawai, "full_name_2") else ""}', set_underline=True, align='center', left_indent=80)
        add_content(doc=doc, content=f'NIP/NRK: {layanan_cuti.pegawai.profil_user.nip if layanan_cuti is not None and hasattr(layanan_cuti.pegawai, "profil_user") and hasattr(layanan_cuti.pegawai.profil_user, "nip") else ""}', left_indent=80, align='center', space_after=20)

        table2 = doc.add_table(rows=4, cols=2)
        self.set_col_widths(table2)
        table2.cell(0,0).merge(table2.cell(0,2))
        add_table_content2(table=table2, row_index=0, col_index=0 , content='CATATAN KEPEGAWAIAN  :\nCuti yang telah diambil dalam tahun yang bersangkutan :\n1. Cuti Tahunan \t: \n2. Cuti Besar  \t\t: \n3. Cuti Sakit \t\t: \n4. Cuti Bersalin \t: \n5. Cuti Karena Alasan Penting  : \n6. Keterangan lain – lain :')
        add_table_content2(table=table2, row_index=0, col_index=1, content='CATATAN PERTIMBANGAN ATASAN LANGSUNG  :', space_after=10)
        add_table_content2(table=table2, row_index=1, col_index=1, content=f'{data_instansi.jabatan_atasan["jabatan_atasan1"] if data_instansi is not None else ""} {data_instansi.jabatan_atasan["instansi1"] if data_instansi is not None else ""} \n')
        table2.cell(row_idx=1, col_idx=1).paragraphs[0].add_run().add_picture(kasi_buffer, Inches(1))
        add_table_content2(table=table2, row_index=1, col_index=1, content=f'\n {data_instansi.nama_atasan["nama_atasan1"] if data_instansi is not None else ""}', set_bold=True, set_underline=True)
        add_table_content2(table=table2, row_index=1, col_index=1, content=f'\nNIP. {data_instansi.nama_atasan["nip_atasan1"] if data_instansi is not None else ""}', set_bold=True, align='center')
        add_table_content2(table=table2, row_index=2, col_index=1, content='MENGETAHUI  :', space_before=10, space_after=6)
        add_table_content2(table=table2, row_index=3, col_index=1, content=f'{data_instansi.jabatan_atasan["jabatan_atasan2"] if data_instansi is not None else ""} {data_instansi.jabatan_atasan["instansi2"] if data_instansi is not None else ""} \n')
        table2.cell(row_idx=3, col_idx=1).paragraphs[0].add_run().add_picture(kabid_buffer, Inches(1))
        add_table_content2(table=table2, row_index=3, col_index=1, content=f'\n {data_instansi.nama_atasan["nama_atasan2"] if data_instansi is not None else ""} ', set_bold=True, set_underline=True)
        add_table_content2(table=table2, row_index=3, col_index=1, style=None, content=f'\nNIP. {data_instansi.nama_atasan["nip_atasan2"] if data_instansi is not None else ""}', set_bold=True, align='center')

        section = doc.sections[0]
        footer = section.footer.paragraphs[0]
        footer.style.font.size=Pt(10.0)
        data_footer = footer.add_run()
        data_footer.add_text("Tanda tangan elektronik ini dapat dilacak keabsahannya melalui link yang terdapat dalam qrcode")
        data_footer.italic=True

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=usulan-cuti-{layanan_cuti.pegawai.first_name}.docx'
        doc.save(response)

        return response


def _safe_get(obj, attr, default="-"):
    try:
        value = getattr(obj, attr, default)
        if value in (None, ""):
            return default
        return value
    except Exception:
        return default


def _status_checkbox_line(keputusan) -> str:
    """
    Untuk bagian VIII:
    [ ] DISETUJUI [ ] PERUBAHAN [ ] DITANGGUHKAN [ ] TIDAK DISETUJUI

    Mendukung nilai workflow baru: belum, setuju, tunda, dan tolak.
    Boolean lama tetap dikenali untuk kompatibilitas data terdahulu.
    """
    labels = ["DISETUJUI", "PERUBAHAN", "DITANGGUHKAN", "TIDAK DISETUJUI"]

    active_index = {
        'setuju': 0,
        'perubahan': 1,
        'tunda': 2,
        'tolak': 3,
        True: 0,
        False: 3,
    }.get(keputusan)

    parts = []
    for i, label in enumerate(labels):
        mark = "✓" if active_index == i else " "
        parts.append(f"[{mark}] {label}")
    return " ".join(parts)


def _build_catatan_cuti_tahunan(
    pegawai,
    tahun_ref: int,
    snapshot=None,
    pada=None,
):
    """Gunakan snapshot saat pengajuan; hitung dinamis hanya untuk data lama."""
    snapshot = snapshot or CheckCuti().buat_snapshot_saldo_cuti(
        pegawai,
        tahun_ref,
        pada=pada,
    )
    rows = []
    for row in snapshot.get('rows', []):
        label = row.get('label', '-')
        tahun = row.get('tahun', '-')
        terpakai = int(row.get('terpakai', 0) or 0)
        sisa_hak = int(row.get('sisa_hak', 0) or 0)
        dapat_digunakan = int(row.get('dapat_digunakan', 0) or 0)
        dicadangkan_default = (
            max(0, sisa_hak - dapat_digunakan)
            if label == 'N'
            else 0
        )
        dicadangkan = int(
            row.get('dicadangkan', dicadangkan_default) or 0
        )
        # Snapshot versi 2 menyimpan cadangan terpisah. Pada formulir, sesuai
        # format ringkas yang diminta, cadangan digabung ke kolom Terpakai.
        if dicadangkan and int(snapshot.get('versi', 1) or 1) <= 2:
            terpakai += dicadangkan
            sisa_hak = max(0, sisa_hak - dicadangkan)
        sisa_tunda = int(row.get('sisa_tunda', 0) or 0)
        # Snapshot versi awal belum menyimpan hak_tunda/terpakai_tunda.
        hak_tunda = int(row.get('hak_tunda', sisa_tunda) or 0)
        terpakai_tunda = int(row.get('terpakai_tunda', 0) or 0)
        if hak_tunda or terpakai_tunda:
            keterangan = (
                f"Hak tunda {hak_tunda} hari; dipakai "
                f"{terpakai_tunda} hari; sisa {sisa_tunda} hari"
            )
        elif label in ('N-1', 'N-2') and dapat_digunakan:
            keterangan = (
                f"Kompensasi hak tahun {tahun}: "
                f"dapat digunakan {dapat_digunakan} hari"
            )
        elif label == 'N':
            keterangan = (
                f"Hak tahun berjalan; terpakai {terpakai} hari; "
                f"sisa {dapat_digunakan} hari"
            )
        else:
            keterangan = "Tidak ada saldo yang dapat digunakan"
        rows.append([
            f"Cuti Tahunan ({label})",
            str(tahun),
            str(terpakai),
            str(sisa_hak),
            str(dapat_digunakan),
            keterangan,
        ])
    return {
        'rows': rows,
        'total_tersedia': int(snapshot.get('total_tersedia', 0) or 0),
        'dibuat_pada': snapshot.get('dibuat_pada', '-'),
    }


def make_qr_drawing(data_str: str, size_mm: float = 25 * mm) -> Drawing:
    """
    Membuat objek Drawing berisi QRCode yang bisa dimasukkan ke story Platypus.
    data_str: string (misal JSON) yang akan dikodekan.
    size_mm: ukuran sisi QR dalam mm.
    """
    qr_code = qr.QrCodeWidget(data_str)
    bounds = qr_code.getBounds()
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]

    # Skala QR agar pas di kotak size_mm x size_mm
    drawing = Drawing(
        size_mm,
        size_mm,
        transform=[size_mm / width, 0, 0, size_mm / height, 0, 0],
    )
    drawing.add(qr_code)
    return drawing

def checkbox_text(label: str, active_label: str) -> str:
    """Teks dengan kotak cek [✓]/[ ] untuk jenis cuti."""
    mark = "✓" if label.lower() == (active_label or "").lower() else " "
    return f"[{mark}] {label}"

def build_header_surat(lokasi_surat, tanggal: date, jabatan_tujuan: str, styles):
    """
    Header surat di pojok kanan atas:
    Lokasi, tanggal
    Kepada
    Yth. {jabatan_tujuan}
    di
    Tempat
    """
    tgl_str = tanggal.strftime("%d-%m-%Y") if tanggal else "................"

    data = [
        [Paragraph(f"{lokasi_surat}, {tgl_str}", styles["BodySmallRight"])],
        [Paragraph("Kepada", styles["BodySmallRight"])],
        [Paragraph(f"Yth. {jabatan_tujuan}", styles["BodySmallRight"])],
        [Paragraph("di", styles["BodySmallRight"])],
        [Paragraph("Tempat", styles["BodySmallRight"])],  # bisa diganti "RSUD Mandalika"
    ]

    t = Table(
        data,
        colWidths=[180 * mm],   # lebar satu kolom, tetap di kanan karena teks right-align
        hAlign="RIGHT",
    )
    t.setStyle(
        TableStyle(
            [
                # tidak pakai BOX / GRID agar seperti contoh BKN (tanpa garis)
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )
    return t

def build_section_I(nama, nip, jabatan, masa_kerja, unit_kerja, styles):
    data = [
        [Paragraph("I. DATA PEGAWAI", styles["BodySmall"]), "", "", ""],
        [
            Paragraph("Nama", styles["BodySmall"]),
            Paragraph(f": {nama}", styles["BodySmall"]),
            Paragraph("NIP", styles["BodySmall"]),
            Paragraph(f": {nip}", styles["BodySmall"]),
        ],
        [
            Paragraph("Jabatan", styles["BodySmall"]),
            Paragraph(f": {jabatan}", styles["BodySmall"]),
            Paragraph("Masa Kerja", styles["BodySmall"]),
            Paragraph(f": {masa_kerja}", styles["BodySmall"]),
        ],
        [
            Paragraph("Unit Kerja", styles["BodySmall"]),
            Paragraph(f": {unit_kerja}", styles["BodySmall"]),
            "","",
        ],
    ]
    t = Table(
        data,
        colWidths=[25 * mm, 70 * mm, 25 * mm, 60 * mm],
        hAlign="LEFT",
    )
    t.setStyle(
        TableStyle(
            [
                ("SPAN", (0, 0), (-1, 0)),  # judul I. DATA PEGAWAI
                ("SPAN", (1, 3), (3, 3)),  # "Unit Kerja" + titik dua gabung
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )
    return t


def build_section_II(jenis_aktif, styles):
    jenis_normal = {
        'Cuti Alasan Penting': 'Cuti Karena Alasan Penting',
        'Cuti Diluar Tanggungan Negara': 'Cuti di Luar Tanggungan Negara',
        'Cuti melahirkan': 'Cuti Melahirkan',
    }.get(jenis_aktif, jenis_aktif)
    data = [
        [Paragraph("II. JENIS CUTI YANG DIAMBIL**", styles["BodySmall"]), "",],
        [
            Paragraph(checkbox_text("Cuti Tahunan", jenis_normal), styles["BodySmall"]),
            Paragraph(checkbox_text("Cuti Besar", jenis_normal), styles["BodySmall"]),
        ],
        [
            Paragraph(checkbox_text("Cuti Sakit", jenis_normal), styles["BodySmall"]),
            Paragraph(checkbox_text("Cuti Melahirkan", jenis_normal), styles["BodySmall"]),
        ],
        [
            Paragraph(
                checkbox_text("Cuti Karena Alasan Penting", jenis_normal),
                styles["BodySmall"],
            ),
            Paragraph(
                checkbox_text("Cuti di Luar Tanggungan Negara", jenis_normal),
                styles["BodySmall"],
            ),
        ],
    ]
    t = Table(
        data,
        colWidths=[90 * mm, 90 * mm],
        hAlign="LEFT",
    )
    t.setStyle(
        TableStyle(
            [
                ("SPAN", (0, 0), (-1, 0)),  # judul II
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )
    return t


def build_section_III(alasan_cuti, styles):
    # Tambah satu baris kosong supaya ada ruang tulis tangan kalau dicetak
    data = [
        [Paragraph("III. ALASAN CUTI", styles["BodySmall"])],
        [Paragraph(alasan_cuti or " ", styles["BodySmall"])],
        [Paragraph(" ", styles["BodySmall"])],
    ]
    t = Table(
        data,
        colWidths=[180 * mm],
        hAlign="LEFT",
    )
    t.setStyle(
        TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("SPAN", (0, 0), (0, 0)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )
    return t


def build_section_IV(lama_hari, tgl_mulai, tgl_akhir, styles):
    tgl_mulai_str = tgl_mulai.strftime("%d-%m-%Y") if tgl_mulai else "......"
    tgl_akhir_str = tgl_akhir.strftime("%d-%m-%Y") if tgl_akhir else "......"

    data = [
        [Paragraph("IV. LAMANYA CUTI", styles["BodySmall"]), "", "", ""],
        [
            Paragraph("Selama", styles["BodySmall"]),
            Paragraph(f"{lama_hari} ( {lama_hari} ) hari", styles["BodySmall"]),
            Paragraph("mulai tanggal", styles["BodySmall"]),
            Paragraph(f"{tgl_mulai_str} s/d {tgl_akhir_str}", styles["BodySmall"]),
        ],
    ]
    t = Table(
        data,
        colWidths=[25 * mm, 55 * mm, 35 * mm, 65 * mm],
        hAlign="LEFT",
    )
    t.setStyle(
        TableStyle(
            [
                ("SPAN", (0, 0), (-1, 0)),
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )
    return t


def build_section_V(catatan_snapshot, styles):
    catatan_rows = catatan_snapshot['rows']
    # catatan_rows sudah berisi 3 baris: N-2, N-1, N
    header1 = [
        Paragraph("V. CATATAN CUTI***", styles["BodySmall"]),
        "", "", "", "", "",
    ]
    header2 = [
        Paragraph("1. CUTI TAHUNAN", styles["BodySmall"]),
        Paragraph("Tahun", styles["BodySmall"]),
        Paragraph("Terpakai", styles["BodySmall"]),
        Paragraph("Sisa Hak", styles["BodySmall"]),
        Paragraph("Dapat Digunakan", styles["BodySmall"]),
        Paragraph("Keterangan", styles["BodySmall"]),
    ]
    body = []
    for label, tahun, terpakai, sisa, komp, keterangan in catatan_rows:
        body.append(
            [
                Paragraph(label, styles["BodySmall"]),
                Paragraph(tahun, styles["BodySmall"]),
                Paragraph(terpakai, styles["BodySmall"]),
                Paragraph(sisa, styles["BodySmall"]),
                Paragraph(komp, styles["BodySmall"]),
                Paragraph(keterangan, styles["BodySmall"]),
            ]
        )

    summary = [
        Paragraph(
            "Total saldo tersedia setelah pengajuan ",
            # f"(snapshot {catatan_snapshot['dibuat_pada']})",
            styles["BodySmall"],
        ),
        "",
        "",
        "",
        Paragraph(
            f"{catatan_snapshot['total_tersedia']} hari",
            styles["BodySmall"],
        ),
        "",
    ]
    data = [header1, header2] + body + [summary]
    summary_row = len(data) - 1

    t = Table(
        data,
        colWidths=[45 * mm, 17 * mm, 20 * mm, 20 * mm, 28 * mm, 50 * mm],
        hAlign="LEFT",
    )
    t.setStyle(
        TableStyle(
            [
                ("SPAN", (0, 0), (-1, 0)),  # judul V
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("BACKGROUND", (0, 1), (-1, 1), colors.lightgrey),
                ("SPAN", (0, summary_row), (3, summary_row)),
                ("SPAN", (4, summary_row), (5, summary_row)),
                ("BACKGROUND", (0, summary_row), (-1, summary_row), colors.whitesmoke),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]
        )
    )
    return t


@login_required
def formulir_cuti_pdf(request, pk):
    """
    Generate PDF 'Formulir Permintaan dan Pemberian Cuti' (Lampiran 1.b)
    untuk satu record RiwayatCuti.
    """
    riwayat = get_object_or_404(
        RiwayatCuti.objects.select_related("pegawai", "usulan"),
        pk=pk,
    )
    pegawai = riwayat.pegawai
    layanan = riwayat.usulan
    if layanan is None or not can_view_leave(request.user, layanan):
        raise PermissionDenied("Anda tidak berhak mengunduh dokumen cuti ini.")

    # URL dokumen (link ke formulir ini sendiri atau ke halaman detail yang Anda mau)
    doc_url = request.build_absolute_uri(request.path)
    # kalau mau ke halaman detail usulan, bisa ganti misal:
    # doc_url = request.build_absolute_uri(
    #     reverse("layanan_cuti_detail", args=[layanan.pk])
    # )

    # ---- Data pegawai / profil (silakan sesuaikan dengan project Anda) ----
    profil = getattr(pegawai, "profil_user", None)
    nama = _safe_get(pegawai, "full_name", _safe_get(pegawai, "first_name", "-"))
    nip = _safe_get(profil, "nip", "-")          # NIP/NRK
    riwayat_jabatan = (
        RiwayatJabatan.objects
        .filter(pegawai=pegawai)
        .select_related('nama_jabatan', 'jenjang_jabatan')
        .order_by('-tmt_jabatan', '-updated_at', '-id')
        .first()
    )
    bagian_jabatan = []
    if riwayat_jabatan:
        if riwayat_jabatan.nama_jabatan:
            bagian_jabatan.append(str(riwayat_jabatan.nama_jabatan))
        if riwayat_jabatan.jenjang_jabatan:
            bagian_jabatan.append(str(riwayat_jabatan.jenjang_jabatan))
        if riwayat_jabatan.detail_nama_jabatan:
            bagian_jabatan.append(riwayat_jabatan.detail_nama_jabatan)
    jabatan = ' - '.join(bagian_jabatan) or '-'

    tanggal_dokumen = (
        layanan.created_at.date()
        if layanan and layanan.created_at
        else date.today()
    )
    pengangkatan = (
        RiwayatPengangkatan.objects
        .filter(pegawai=pegawai, tmt_pegawai__isnull=False)
        .order_by('-tmt_pegawai', '-id')
        .first()
    )
    masa_kerja = '-'
    if pengangkatan and pengangkatan.tmt_pegawai:
        masa = relativedelta(tanggal_dokumen, pengangkatan.tmt_pegawai)
        masa_kerja = f'{max(0, masa.years)} tahun {max(0, masa.months)} bulan'

    # Riwayat penempatan aktif
    penempatan = (
        RiwayatPenempatan.objects.filter(pegawai=pegawai, status=True)
        .select_related(
            "penempatan_level1__satker_induk__instansi_daerah",
            "penempatan_level2__unor__satker_induk__instansi_daerah",
            "penempatan_level3__bidang__unor__satker_induk__instansi_daerah",
            "penempatan_level4__sub_bidang__bidang__unor__satker_induk__instansi_daerah",
        )
        .order_by('-updated_at', '-id')
        .first()
    )
    unit_kerja = penempatan.penempatan if penempatan else "-"
    # _, level = penempatan._penempatan_aktif
    
    # Alamat & telp saat cuti
    alamat_cuti = riwayat.domisili_saat_cuti or "-"
    telp = _safe_get(profil, "no_hp", "-")  # sesuaikan field no HP/telepon

    # Tahun referensi (N)
    tahun_ref = riwayat.tahun_cuti or date.today().year

    # Verifikasi cuti (3 level)
    verifikasi = getattr(layanan, "verifikasicuti", None)
    if verifikasi is None and layanan is not None:
        try:
            verifikasi = VerifikasiCuti.objects.get(layanan_cuti=layanan)
        except VerifikasiCuti.DoesNotExist:
            verifikasi = None

    # Gabungkan label struktur aktif dengan identitas verifikator yang dibekukan.
    chain_aktif = {
        item['level']: item
        for item in build_approval_chain(pegawai)
    }
    chain = []
    for level in (1, 2, 3):
        item = chain_aktif.get(level)
        saved_user = (
            getattr(verifikasi, f'verifikator{level}', None)
            if verifikasi else None
        )
        if item is None and saved_user is None:
            continue
        active_user = item['user'] if item else None
        label = (
            item['label']
            if item and (
                saved_user is None
                or active_user is None
                or active_user.pk == saved_user.pk
            )
            else f'Verifikator Level {level}'
        )
        chain.append({
            'level': level,
            'user': saved_user or active_user,
            'label': label,
            'keputusan': (
                getattr(verifikasi, f'keputusan{level}', 'belum')
                if verifikasi else 'belum'
            ),
        })

    direct_entry = chain[0] if chain else None
    final_entry = chain[-1] if chain else None
    user_atasan = direct_entry['user'] if direct_entry else None
    jabatan_atasan = (
        direct_entry['label']
        if direct_entry else 'Atasan Langsung'
    )

    # Nama & NIP atasan diambil langsung dari user_atasan (sesuai TTE)
    nama_atasan2 = "-"
    nip_atasan2 = "-"
    if user_atasan:
        nama_atasan2 = _safe_get(user_atasan, "full_name_2",
                                 _safe_get(user_atasan, "full_name", "-"))
        profil_atasan = getattr(user_atasan, "profil_user", None)
        nip_atasan2 = _safe_get(profil_atasan, "nip", "-")

    status_line_lvl1 = _status_checkbox_line(
        direct_entry['keputusan'] if direct_entry else 'belum'
    )
    
    # ====== Pejabat Penandatangan TERTINGGI (final signer) ======
    user_atasan3 = final_entry['user'] if final_entry else None
    final_label = (
        final_entry['label']
        if final_entry else 'Pejabat Berwenang'
    )


    nama_pejabat = "-"
    nip_pejabat = "-"
    if user_atasan3:
        nama_pejabat = _safe_get(user_atasan3, "full_name_2",
                                 _safe_get(user_atasan3, "full_name", "-"))
        profil_pejabat = getattr(user_atasan3, "profil_user", None)
        nip_pejabat = _safe_get(profil_pejabat, "nip", "-")

    jabatan_pejabat = final_label

    status_line_lvl3 = _status_checkbox_line(
        final_entry['keputusan'] if final_entry else 'belum'
    )

    # ---------------- QR DATA (TTD ELEKTRONIK) ----------------
    qr_data_pemohon = {
        "jenis": "TTE Pemohon Cuti",
        "nama": nama,
        "nip": nip,
        "dokumen": doc_url,
    }
    qr_data_atasan = {
        "jenis": "TTE Atasan Langsung",
        "nama": nama_atasan2,
        "nip": nip_atasan2,
        "dokumen": doc_url,
    }
    qr_data_pejabat = {
        "jenis": "TTE Pejabat Berwenang",
        "nama": nama_pejabat,
        "nip": nip_pejabat,
        "dokumen": doc_url,
    }

    qr_str_pemohon = json.dumps(qr_data_pemohon, ensure_ascii=False)
    qr_str_atasan = json.dumps(qr_data_atasan, ensure_ascii=False)
    qr_str_pejabat = json.dumps(qr_data_pejabat, ensure_ascii=False)

    qr_pemohon_drawing = make_qr_drawing(qr_str_pemohon, size_mm=25 * mm)
    qr_atasan_drawing = None
    qr_pejabat_drawing = None
    
    if direct_entry and direct_entry['keputusan'] == 'setuju' and user_atasan:
        qr_atasan_drawing = make_qr_drawing(qr_str_atasan, size_mm=25 * mm)
    if final_entry and final_entry['keputusan'] == 'setuju' and user_atasan3:
        qr_pejabat_drawing = make_qr_drawing(qr_str_pejabat, size_mm=25 * mm)

    # ==========================================================
    #               SETUP REPORTLAB PDF
    # ==========================================================
    response = HttpResponse(content_type="application/pdf")
    nama_file = slugify(nama) or f'pegawai-{pegawai.pk}'
    filename = f"Formulir_Cuti_{nama_file}_{tahun_ref}.pdf"
    response["Content-Disposition"] = f'inline; filename="{filename}"'

    buffer = BytesIO()
    FOLIO = (215 * mm, 330 * mm)

    doc = SimpleDocTemplate(
        buffer,
        pagesize=FOLIO,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="HeadingSmall",
            parent=styles["Normal"],
            fontSize=10,
            leading=12,
            spaceAfter=4,
            spaceBefore=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodySmall", parent=styles["Normal"], fontSize=9, leading=11
        )
    )
    
    styles.add(
        ParagraphStyle(
            name="BodySmallRight",
            parent=styles["BodySmall"],
            alignment=TA_RIGHT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodySmallIndent",
            parent=styles["BodySmall"],
            leftIndent=6 * mm,   # kira-kira 1 tab ke dalam
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyMoreIndent",
            parent=styles["BodySmall"],
            leftIndent=12 * mm,   # kira-kira 2 tab ke dalam
        )
    )
    
    styles.add(
        ParagraphStyle(
            name="TitleSmall",
            parent=styles["Title"],
            fontSize=12,      # <--- ukuran ideal
            leading=14,
            alignment=1,      # CENTER
            spaceAfter=6,
            spaceBefore=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodySmallTight",
            parent=styles["BodySmall"],
            leading=9,     # lebih rapat (default biasanya 11–12)
            spaceBefore=0,
            spaceAfter=0,
        )
    )

    elements = []

    # Header surat: lokasi, tanggal, kepada Yth.
    lokasi_surat = "Lombok Tengah"  # bisa Bapak ganti dari settings / InstansiDaerah

    tanggal_str = tanggal_dokumen.strftime("%d-%m-%Y")

    header_rows = [
        ["", Paragraph(f"{lokasi_surat}, {tanggal_str}", styles["BodySmall"])],
        ["", Paragraph("Kepada", styles["BodySmallIndent"])],
        ["", Paragraph(f"Yth. {jabatan_pejabat}", styles["BodySmall"])],
        ["", Paragraph("di -", styles["BodySmallIndent"])],
        ["", Paragraph("Tempat", styles["BodyMoreIndent"])],
    ]

    header_table = Table(
        header_rows,
        colWidths=[100*mm, 80*mm],   # kiri kosong, kanan isi header
        hAlign="RIGHT",
    )
    header_table.setStyle(
        TableStyle([
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("FONTNAME", (0,0), (-1,-1), "Helvetica"),
            ("FONTSIZE", (0,0), (-1,-1), 9),
        ])
    )

    elements.append(header_table)
    elements.append(Spacer(1, 12))

    # Judul formulir
    elements.append(
        Paragraph("FORMULIR PERMINTAAN DAN PEMBERIAN CUTI", styles["TitleSmall"])
    )
    elements.append(Spacer(1, 8))

    # I. DATA PEGAWAI
    elements.append(
        build_section_I(nama, nip, jabatan, masa_kerja, unit_kerja, styles)
    )
    elements.append(Spacer(1, 4))

    # II. JENIS CUTI YANG DIAMBIL
    jenis_aktif = (
        riwayat.get_jenis_cuti_display()
        if hasattr(riwayat, "get_jenis_cuti_display")
        else riwayat.jenis_cuti
    )
    elements.append(build_section_II(jenis_aktif, styles))
    elements.append(Spacer(1, 4))

    # III. ALASAN CUTI
    elements.append(build_section_III(riwayat.alasan_cuti, styles))
    elements.append(Spacer(1, 4))

    # IV. LAMANYA CUTI
    lama_hari = riwayat.lama_cuti or 0
    elements.append(
        build_section_IV(
            lama_hari,
            riwayat.tgl_mulai_cuti,
            riwayat.tgl_akhir_cuti,
            styles,
        )
    )
    elements.append(Spacer(1, 4))

    # V. CATATAN CUTI – pakai CheckCuti
    catatan_snapshot = _build_catatan_cuti_tahunan(
        pegawai,
        tahun_ref,
        snapshot=layanan.snapshot_saldo_cuti,
        pada=tanggal_dokumen,
    )
    elements.append(build_section_V(catatan_snapshot, styles))
    elements.append(Spacer(1, 8))

    # VI. ALAMAT SELAMA MENJALANKAN CUTI
    elements.append(
        Paragraph("VI. ALAMAT SELAMA MENJALANKAN CUTI", styles["HeadingSmall"])
    )
    alamat_rows = [
        ["Alamat", f": {alamat_cuti}"],
        ["Telp", f": {telp}"],
    ]
    t_alamat = Table(alamat_rows, colWidths=[30 * mm, 140 * mm])
    t_alamat.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )
    elements.append(t_alamat)
    elements.append(Spacer(1, 12))

    # Tanda tangan pemohon
    tte_pemohon = [
        # baris 1: tanggal + "Hormat saya," dalam 1 paragraf, 2 line
        [
            "",
            Paragraph(
                f"Lombok Tengah, {tanggal_dokumen.strftime('%d-%m-%Y')}<br/>Hormat saya,",
                styles["BodySmall"],
            ),
        ],
        # baris 2: QR di kanan
        [
            "",
            qr_pemohon_drawing,
        ],
        # baris 3: nama + NIP dalam 1 paragraf, 2 line
        [
            "",
            Paragraph(
                f"{nama}<br/>NIP. {nip}",
                styles["BodySmall"],
            ),
        ],
    ]

    tte_pemohon_qr_table = Table(
        tte_pemohon,
        colWidths=[100 * mm, 80 * mm],  # kiri kosong, kanan isi TTE
        hAlign="RIGHT",
    )
    tte_pemohon_qr_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )

    elements.append(tte_pemohon_qr_table)
    elements.append(Spacer(1, 6))
    
    # VII. PERTIMBANGAN ATASAN LANGSUNG (dengan TTE QR)
    elements.append(
        Paragraph("VII. PERTIMBANGAN ATASAN LANGSUNG", styles["HeadingSmall"])
    )
    elements.append(Paragraph(status_line_lvl1, styles["BodySmall"]))
    elements.append(Spacer(1, 8))
    
    ttd_atasan_table = [
        [
            "",
            qr_atasan_drawing,
        ],
        [
            "",
            Paragraph(
                f"{jabatan_atasan}<br/>{nama_atasan2}<br/>NIP. {nip_atasan2}",
                styles["BodySmall"],
            ),
        ],
    ]

    ttd_atasan_table_qr_table = Table(
        ttd_atasan_table,
        colWidths=[100 * mm, 80 * mm],  # kiri kosong, kanan isi TTE
        hAlign="RIGHT",
    )
    ttd_atasan_table_qr_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )

    elements.append(ttd_atasan_table_qr_table)
    elements.append(Spacer(1, 12))

    # VIII. KEPUTUSAN PEJABAT YANG BERWENANG MEMBERIKAN CUTI (dengan TTE QR)
    elements.append(
        Paragraph(
            "VIII. KEPUTUSAN PEJABAT YANG BERWENANG MEMBERIKAN CUTI",
            styles["HeadingSmall"],
        )
    )
    elements.append(Paragraph(status_line_lvl3, styles["BodySmall"]))
    elements.append(Spacer(1, 8))
    
    ttd_pejabat_table = [
        [
            "",
            qr_pejabat_drawing,
        ],
        [
            "",
            Paragraph(
                f"{jabatan_pejabat}<br/>{nama_pejabat}<br/>NIP. {nip_pejabat}",
                styles["BodySmall"],
            ),
        ],
    ]

    ttd_pejabat_table_qr_table = Table(
        ttd_pejabat_table,
        colWidths=[100 * mm, 80 * mm],  # kiri kosong, kanan isi TTE
        hAlign="RIGHT",
    )
    ttd_pejabat_table_qr_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )

    elements.append(ttd_pejabat_table_qr_table)

    # Build PDF
    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()
    response.write(pdf)
    return response


class LayananCutiDocxView(LoginRequiredMixin, View):
    def set_col_widths(self, table):
        widths = (Inches(0.3), Inches(7.7))
        # height = (Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                # row.cells[idx].height = heigh

    def get_object(self, id):
        try:
            data = LayananCuti.objects.get(id=id)
            return data
        except Exception:
            return None
        
    def get_riwayat_cuti(self, id):
        try:
            data = RiwayatCuti.objects.filter(usulan__id=id).last()
            return data
        except RiwayatCuti.DoesNotExist:
            return None
        
    def get_data_cuti(self, data_input=None, data_output=None, attr=""):
        if data_input is not None and hasattr(data_input.cuti_usulan, attr):
            return data_output
        return ""
    
    def get_unor_pimpinan(self, pegawai, layanan_cuti=None):
        try:
            penempatan = RiwayatPenempatan.objects.filter(pegawai=pegawai, status=True).order_by('-id').first()
            if penempatan:
                data = penempatan.unor_pimpinan
                snapshot = getattr(layanan_cuti, 'verifikasicuti', None)
                if snapshot:
                    signer = next((user for user in (
                        snapshot.verifikator3,
                        snapshot.verifikator2,
                        snapshot.verifikator1,
                    ) if user is not None), None)
                    if signer:
                        data['nama_pimpinan'] = getattr(signer, 'full_name_2', 'N/A')
                        profil = getattr(signer, 'profil_user', None)
                        data['nip'] = getattr(profil, 'nip', 'N/A') or 'N/A'
                        panggol = signer.riwayatpanggol_set.order_by('-id').first()
                        data['panggol'] = getattr(panggol, 'panggol', 'N/A') or 'N/A'
                        struktur = penempatan.unor
                        masa_jabatan = struktur.riwayat_pejabat.filter(
                            pejabat=signer,
                        ).order_by('-tanggal_mulai', '-id').first() if struktur else None
                        if masa_jabatan and masa_jabatan.nama_jabatan:
                            data['pimpinan'] = masa_jabatan.nama_jabatan
                return data
            return ""
        except Exception:
            return ""
        
    def get_status_pegawai(self, pegawai):
        try:
            pengangkatan = RiwayatPengangkatan.objects.filter(pegawai=pegawai).order_by('-id').first()
            if pengangkatan is None:
                return None
            return pengangkatan.desk_status_pegawai
        except RiwayatPengangkatan.DoesNotExist:
            return None

    def get(self, request, **kwargs):
        #Pengaturan ukuran kertas
        doc:Document=CreateDocument()
        id = kwargs.get('layanan_id')
        layanan_cuti = self.get_object(id=id)
        if layanan_cuti is None or not can_view_leave(request.user, layanan_cuti):
            raise PermissionDenied("Anda tidak berhak mengunduh dokumen cuti ini.")
        riwayat_cuti = self.get_riwayat_cuti(id=id)
        jabatan = riwayat_cuti.pegawai.riwayatjabatan_set.last()
        data_instansi = riwayat_cuti.pegawai.riwayat_penempatan.last()
        lama_cuti = riwayat_cuti.lama_cuti if hasattr(riwayat_cuti, "lama_cuti") else 0
        riwayatpenempatan = riwayat_cuti.pegawai.riwayat_penempatan.last()
        page_size = doc.sections[0]
        page_size.page_width = Mm(215.9)
        page_size.page_height = Mm(330.0)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.bottom_margin = Inches(1.0)
        #penambahan gambar kop
        doc.paragraphs.clear()
        doc.add_picture('static/img/KOP RS MANDALIKA 2024.png', width=Inches(7.0))
        add_content(doc=doc, content2=f'SURAT IZIN {riwayat_cuti.jenis_cuti.upper() if hasattr(riwayat_cuti, "jenis_cuti") else ""}', align='center', set_bold=True, set_underline=True)
        add_content(doc=doc, content2='Nomor : ${nomor_naskah}', align='center')
        # add_content(doc=doc, 'Lombok Tengah, ${tanggal_naskah}', align='right', space_after=6)
        
        add_content(doc=doc, content=f'Diberikan {riwayat_cuti.jenis_cuti.lower() if hasattr(riwayat_cuti, "jenis_cuti") else ""} kepada {self.get_status_pegawai(riwayat_cuti.pegawai) if self.get_status_pegawai(riwayat_cuti.pegawai) else ""},', space_before=10, space_after=10, style='List Number')
        add_content(doc=doc, content=f'nama \t\t: {riwayat_cuti.pegawai.full_name_2 if hasattr(riwayat_cuti.pegawai, "full_name_2") else ""},', tab=1)
        add_content(doc=doc, content=f'NIP/NRK \t\t: {riwayat_cuti.pegawai.profil_user.nip if hasattr(riwayat_cuti.pegawai, "profil_user") and hasattr(riwayat_cuti.pegawai.profil_user, "nip") else ""},', tab=1)
        add_content(doc=doc, content=f'jabatan \t\t: {jabatan.nama_jabatan if jabatan is not None and hasattr(jabatan, "nama_jabatan") else ""}', tab=1)
        add_content(doc=doc, content=f'satuan organisasi \t: {data_instansi.unor if data_instansi is not None else "" },', tab=1)
        add_content(doc=doc, content=f'selama {convert_num_2_word(lama_cuti)} ({lama_cuti}) hari, terhitung mulai tanggal {riwayat_cuti.tgl_mulai_cuti.strftime("%d %B %Y") if hasattr(riwayat_cuti, "tgl_mulai_cuti") and hasattr(riwayat_cuti.tgl_mulai_cuti, "strftime") else ""} s.d {riwayat_cuti.tgl_akhir_cuti.strftime("%d %B %Y") if hasattr(riwayat_cuti ,"tgl_akhir_cuti") and hasattr(riwayat_cuti.tgl_akhir_cuti, "strftime") else ""} dengan ketentuan sebagai berikut :', space_after=10, space_before=10)
        
        table2 = doc.add_table(rows=2, cols=2)
        self.set_col_widths(table2)
        add_table_content2(table=table2, row_index=0, col_index=0, content='a. ')
        add_table_content2(table=table2, row_index=0, col_index=1, content=f'Sebelum menjalankan {riwayat_cuti.jenis_cuti.lower() if hasattr(riwayat_cuti, "jenis_cuti") else ""} wajib menyerahkan pekerjaannya kepada atasan langsungnya atau pejabat lain yang ditentukan.')
        add_table_content2(table=table2, row_index=1, col_index=0, content='b. ')
        add_table_content2(table=table2, row_index=1, col_index=1, style=None, content=f'Setelah selesai menjalankan {riwayat_cuti.jenis_cuti.lower() if hasattr(riwayat_cuti, "jenis_cuti") else ""} wajib melaporkan diri kepada atasan langsungnya dan bekerja kembali sebagaimana mestinya.')
        
        pimpinan = riwayatpenempatan.unor_pimpinan if riwayatpenempatan is not None else ""
        add_content(doc=doc, content='Demikian surat cuti ini dibuat untuk dapat dipergunakan sebagaimana mestinya.', space_before=10, space_after=16, style='List Number')
        add_content(doc=doc, content='Lombok Tengah, ${tanggal_naskah}', align='left', left_indent=80)
        add_content(doc=doc, content=f'{self.get_unor_pimpinan(layanan_cuti.pegawai, layanan_cuti)["pimpinan"] if self.get_unor_pimpinan(layanan_cuti.pegawai, layanan_cuti) else ""}', left_indent=80, align='left', space_after=30)
        add_content(doc=doc, content='${ttd_pengirim}', left_indent=80, align='left', space_after=30)
        add_content(doc=doc, content=f'{self.get_unor_pimpinan(layanan_cuti.pegawai, layanan_cuti)["nama_pimpinan"] if self.get_unor_pimpinan(layanan_cuti.pegawai, layanan_cuti) else ""}', left_indent=80, align='left')
        add_content(doc=doc, content=f'{self.get_unor_pimpinan(layanan_cuti.pegawai, layanan_cuti)["panggol"] if self.get_unor_pimpinan(layanan_cuti.pegawai, layanan_cuti) else ""}', left_indent=80, align='left')
        add_content(doc=doc, content=f'{self.get_unor_pimpinan(layanan_cuti.pegawai, layanan_cuti)["nip"] if self.get_unor_pimpinan(layanan_cuti.pegawai, layanan_cuti) else ""}', left_indent=80, align='left')

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=cuti-{riwayat_cuti.pegawai.full_name}.docx'
        doc.save(response)

        return response

###################################### PELIMPAHAN TUGAS PDF #########################################
def _get_nip(user) -> str:
    profil = getattr(user, "profil_user", None)
    return getattr(profil, "nip", None) or "-"


def _get_rp_aktif(user):
    """
    Ambil riwayat penempatan aktif terbaru.
    """
    if not user:
        return None
    qs = getattr(user, "riwayat_penempatan", None)
    if not qs:
        return None
    return qs.filter(status=True).order_by("-updated_at", "-id").first()


def _resolve_kepala_instalasi(pelimpahan: PelimpahanTugas):
    """
    Tentukan kepala instalasi untuk blok tanda tangan.
    Prioritas:
    1) pelimpahan.kepala_instalasi (kalau sudah diputuskan)
    2) dari riwayat pejabat aktif pada UnitInstalasi penempatan pemberi tugas
    """
    if pelimpahan.kepala_instalasi:
        return pelimpahan.kepala_instalasi

    rp = _get_rp_aktif(pelimpahan.pemberi_tugas)
    if not rp:
        return None

    obj, level = rp._penempatan_aktif
    if level == "level4" and obj:
        return get_active_leader(obj)

    # jika bukan level4, tidak dipaksakan (Anda bisa sesuaikan bila ingin kepala unit lain)
    return None


def _qr_image(payload: str, size_cm: float = 2.5) -> Image:
    qr = qrcode.QRCode(
        version=None,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=10,
        border=2,
    )
    qr.add_data(payload)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")

    bio = BytesIO()
    img.save(bio, format="PNG")
    bio.seek(0)
    return Image(bio, width=size_cm * cm, height=size_cm * cm)


@login_required
def pelimpahan_tugas_pdf(request, pk: int):
    pelimpahan = get_object_or_404(
        PelimpahanTugas.objects.select_related(
            "riwayat_cuti", "pemberi_tugas", "penerima_tugas", "atasan_penyetuju"
        ),
        pk=pk
    )
    if not can_view_delegation(request.user, pelimpahan):
        raise PermissionDenied("Anda tidak berhak mengunduh dokumen pelimpahan ini.")

    pdf_url = request.build_absolute_uri(
        reverse("file_urls:pelimpahan_tugas_pdf", kwargs={"pk": pelimpahan.pk})
    )

    pemberi = pelimpahan.pemberi_tugas
    penerima = pelimpahan.penerima_tugas

    butuh_atasan = pelimpahan.requires_atasan_approval()
    atasan = pelimpahan.atasan_penyetuju if butuh_atasan else None

    # ---- flag tampil QR ----
    show_qr_penerima = (pelimpahan.persetujuan_penerima == "disetujui")
    show_qr_atasan = (butuh_atasan and pelimpahan.persetujuan_atasan == "disetujui")

    jab_pemberi, unit_pemberi = get_jabatan_unit(pemberi, mode="top")   # atau "full"
    jab_penerima, unit_penerima = get_jabatan_unit(penerima, mode="top")

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2.2*cm, rightMargin=2.2*cm,
        topMargin=2.0*cm, bottomMargin=2.0*cm
    )

    styles = getSampleStyleSheet()

    title = ParagraphStyle(
        "title",
        parent=styles["Title"],
        alignment=1,
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
    )

    normal = ParagraphStyle(
        "normal",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=15,
    )

    normal_center = ParagraphStyle(
        "normal_center",
        parent=normal,
        alignment=1,
    )

    small = ParagraphStyle(
        "small",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
    )

    small_center = ParagraphStyle(
        "small_center",
        parent=small,
        alignment=1,
    )

    story = []
    story.append(Paragraph("<u>SURAT PELIMPAHAN TUGAS</u>", title))
    story.append(Spacer(1, 0.6 * cm))

    deskripsi = (pelimpahan.deskripsi_tugas or "").replace("\n", "<br/>")

    body = f"""
    Yang bertanda tangan di bawah ini :<br/><br/>

    Nama&nbsp;&nbsp;&nbsp;: {getattr(pemberi, "full_name", str(pemberi))}<br/>
    NIP&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;: {get_nip(pemberi)}<br/>
    Jabatan : {jab_pemberi}<br/>
    Unit&nbsp;&nbsp;&nbsp;&nbsp;: {unit_pemberi}<br/><br/>

    Dengan ini melimpahkan tugas kepada :<br/><br/>

    Nama&nbsp;&nbsp;&nbsp;: {getattr(penerima, "full_name", str(penerima))}<br/>
    NIP&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;: {get_nip(penerima)}<br/>
    Jabatan : {jab_penerima}<br/>
    Unit&nbsp;&nbsp;&nbsp;&nbsp;: {unit_penerima}<br/><br/>

    Untuk menggantikan saya sementara melaksanakan tugas dan fungsi saya selama menjalankan cuti terhitung
    sejak tanggal {pelimpahan.tgl_mulai} s.d {pelimpahan.tgl_selesai}<br/><br/>

    Demikian surat pelimpahan tugas ini dibuat untuk dipergunakan sebagaimana mestinya.
    """
    story.append(Paragraph(body, normal))
    story.append(Spacer(1, 1.0 * cm))

    # ---- QR payload + image ----
    payload_pemberi = (
        f"Nama: {getattr(pemberi,'full_name',str(pemberi))}\n"
        f"NIP: {get_nip(pemberi)}\n"
        f"Link: {pdf_url}"
    )
    qr_pemberi = _qr_image(payload_pemberi)

    payload_penerima = (
        f"Nama: {getattr(penerima,'full_name',str(penerima))}\n"
        f"NIP: {get_nip(penerima)}\n"
        f"Link: {pdf_url}"
    )
    qr_penerima = _qr_image(payload_penerima)

    left = [
        Paragraph("Yang Memberi Mandat,", normal_center),
        Spacer(1, 0.2 * cm),
        qr_pemberi,
        Spacer(1, 0.2 * cm),
        Paragraph(f"<b><u>{getattr(pemberi,'full_name_2', getattr(pemberi,'full_name',str(pemberi)))}</u></b>", small_center),
        Paragraph(f"NIP : {get_nip(pemberi)}", small_center),
    ]

    right = [
        Paragraph("Yang Diberi Mandat,", normal_center),
        Spacer(1, 0.2 * cm),
    ]

    # QR penerima hanya muncul jika penerima setuju
    if show_qr_penerima:
        right += [qr_penerima, Spacer(1, 0.2 * cm)]
    else:
        # spacer biar layout tetap rapi walau QR disembunyikan
        right += [Spacer(1, 3.0 * cm)]

    right += [
        Paragraph(f"<b><u>{getattr(penerima,'full_name_2', getattr(penerima,'full_name',str(penerima)))}</u></b>", small_center),
        Paragraph(f"NIP : {get_nip(penerima)}", small_center),
    ]

    t = Table([[left, right]], colWidths=[8.5 * cm, 8.5 * cm])
    t.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
    ]))
    story.append(t)

    # ---- Blok atasan hanya jika pemberi level4 ----
    if atasan:
        story.append(Spacer(1, 1.2 * cm))

        payload_atasan = (
            f"Nama: {getattr(atasan,'full_name_2', getattr(atasan,'full_name',str(atasan)))}\n"
            f"NIP: {get_nip(atasan)}\n"
            f"Link: {pdf_url}"
        )
        qr_atasan = _qr_image(payload_atasan)

        story.append(Paragraph("Disetujui oleh Kepala Instalasi/Unit", normal_center))
        story.append(Spacer(1, 0.25 * cm))

        if show_qr_atasan:
            story.append(qr_atasan)
            story.append(Spacer(1, 0.2 * cm))
        else:
            story.append(Spacer(1, 3.0 * cm))

        story.append(Paragraph(
            f"<b><u>{getattr(atasan,'full_name_2', getattr(atasan,'full_name',str(atasan)))}</u></b>",
            small_center
        ))
        story.append(Paragraph(f"NIP : {get_nip(atasan)}", small_center))

    doc.build(story)

    pdf = buffer.getvalue()
    buffer.close()

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = f'inline; filename="surat_pelimpahan_tugas_{pelimpahan.pk}.pdf"'
    return response



class FormatPenilaianInovasiView(View):
    def set_col_widths(self, table):
        widths = (Inches(0.3), Inches(4.0), Inches(8.0), Inches(3.0), Inches(3.0))
        # height = (Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                # row.cells[idx].height = height

    def get_object(self, id):
        try:
            data = LayananUsulanInovasi.objects.get(id=id)
            return data
        except Exception:
            return None

    def get(self, request, **kwargs):
        #Pengaturan ukuran kertas
        # clear_document(doc)
        doc:Document=CreateDocument()
        id = kwargs.get('layanan_id')
        layanan_inovasi = self.get_object(id=id)
        page_size = doc.sections[0]
        page_size.page_width = Mm(215.9)
        page_size.page_height = Mm(330.0)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.bottom_margin = Inches(1.0)
        #penambahan gambar kop
        doc.paragraphs.clear()
        doc.add_picture('static/img/KOP RS MANDALIKA 2024.png', width=Inches(7.0))
        add_content(doc=doc, content2=f'FORMAT PENILAIAN DAN PEMBINAAN KATEGORI INOVASI', align='center', set_bold=True, set_underline=True)
        add_content(doc=doc, content2='', align='center', space_after=10)
        # add_content(doc=doc, 'Lombok Tengah, ${tanggal_naskah}', align='right', space_after=6)

        data_penempatan = layanan_inovasi.pegawai.riwayat_penempatan.last().penempatan
        add_content(doc=doc, content=f'Nama Nakes/Named \t\t: {layanan_inovasi.pegawai.full_name if layanan_inovasi is not None and hasattr(layanan_inovasi.pegawai, "full_name") else ""},', tab=1)
        add_content(doc=doc, content=f'Jenis Nakes/Named \t\t: {layanan_inovasi.pegawai.riwayatjabatan_set.last().nama_jabatan if layanan_inovasi is not None and hasattr(layanan_inovasi, "pegawai") else ""},', tab=1)
        add_content(doc=doc, content=f'Instalasi \t\t\t: {data_penempatan}', tab=1)
        add_content(doc=doc, content=f'Judul Inovasi \t\t\t: {layanan_inovasi.inovasi.judul if layanan_inovasi is not None and hasattr(layanan_inovasi.inovasi, "judul") else ""},', tab=1, space_after=10)

        table2 = doc.add_table(rows=6, cols=5)
        self.set_col_widths(table2)
        table2.cell(1,4).merge(table2.cell(4,4))
        table2.cell(5,1).merge(table2.cell(5,2))
        add_table_content2(table=table2, row_index=0, col_index=0, content='No.')
        add_table_content2(table=table2, row_index=0, col_index=1, content='Aspek Penilaian')
        add_table_content2(table=table2, row_index=0, col_index=2, content='Unsur Penilaian')
        add_table_content2(table=table2, row_index=0, col_index=3, content='Nilai/Poin')
        add_table_content2(table=table2, row_index=0, col_index=4, content='Total Nilai')

        add_table_content2(table=table2, row_index=1, col_index=0, content='1.')
        add_table_content2(table=table2, row_index=1, col_index=1, content='Makalah/Essay \n(Skor: 30)')
        add_table_content2(table=table2, row_index=1, col_index=2, content='1. Kriteria Umum (Teknis Penulisan)')
        add_table_content2(table=table2, row_index=1, col_index=3, content='')
        add_table_content2(table=table2, row_index=1, col_index=4, content='')

        add_table_content2(table=table2, row_index=2, col_index=0, content='')
        add_table_content2(table=table2, row_index=2, col_index=1, content='')
        add_table_content2(table=table2, row_index=2, col_index=2, content='2. Kriteria Khusus (Keberhasilan dan Keberlanjutan Program)')
        add_table_content2(table=table2, row_index=2, col_index=3, content='')
        add_table_content2(table=table2, row_index=2, col_index=4, content='')

        add_table_content2(table=table2, row_index=3, col_index=0, content='2.')
        add_table_content2(table=table2, row_index=3, col_index=1, content='Presentasi dan Wawancara \n(Skor: 40)')
        add_table_content2(table=table2, row_index=3, col_index=2, content='1. Materi \n2. Kemampuan Menjawab Pertanyaan \n3. Kemampuan Presentasi')
        add_table_content2(table=table2, row_index=3, col_index=3, content='')
        add_table_content2(table=table2, row_index=3, col_index=4, content='')

        add_table_content2(table=table2, row_index=4, col_index=0, content='3.')
        add_table_content2(table=table2, row_index=4, col_index=1, content='Daya ungkit Inovasi dan Dampak \n(Skor: 30)')
        add_table_content2(table=table2, row_index=4, col_index=2, content='')
        add_table_content2(table=table2, row_index=4, col_index=3, content='')
        add_table_content2(table=table2, row_index=4, col_index=4, content='')

        add_table_content2(table=table2, row_index=5, col_index=0, content='')
        add_table_content2(table=table2, row_index=5, col_index=1, content='TOTAL NILAI')
        add_table_content2(table=table2, row_index=5, col_index=2, content='')
        add_table_content2(table=table2, row_index=5, col_index=3, content='')
        add_table_content2(table=table2, row_index=5, col_index=4, content='')
        
        add_content(doc=doc, content='', space_before=10)
        add_content(doc=doc, content='Lombok Tengah, ${tanggal_naskah}', align='center', left_indent=80)
        add_content(doc=doc, content='${jabatan_pengirim}', left_indent=80, align='center')
        add_content(doc=doc, content='${ttd_pengirim}', left_indent=80, align='center')
        add_content(doc=doc, content='${nama_pengirim}', align='center', left_indent=80)
        add_content(doc=doc, content='NIP. ', left_indent=80, align='center')

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=format-penilaian-inovasi ({layanan_inovasi.pegawai.full_name}).docx'
        doc.save(response)
        return response

#SPT diklat untuk satu orang
class LayananDiklatSPTDocxView(View):
    def set_col_widths(self, table):
        widths = (Inches(2), Inches(0.3), Inches(10))
        # height = (Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                # row.cells[idx].height = heigh

    def get_object(self, id):
        try:
            data = RiwayatDiklat.objects.get(id=id)
            return data
        except Exception:
            return None

    def get(self, request, **kwargs):
        #Pengaturan ukuran kertas
        doc:Document=CreateDocument()
        id = kwargs.get('diklat_id')
        riwayatdiklat = self.get_object(id=id)
        spt_text = TextSPTDiklat.objects.filter(diklat=riwayatdiklat.usulan).first()
        jabatan = None
        panggol_sdm = None
        panggol = None
        penempatan = None
        pegawai = riwayatdiklat.pegawai.first()
        if pegawai:
            jabatan = pegawai.riwayatjabatan_set.last()
            penempatan = RiwayatPenempatan.objects.filter(pegawai=pegawai, status=True).last()
            panggol_sdm = pegawai.riwayatpanggol_set.last()
            panggol = get_active_leader(penempatan.penempatan_level1).riwayatpanggol_set.last() if hasattr(penempatan, 'penempatan_level1') and get_active_leader(penempatan.penempatan_level1) else None
        # Parse the HTML content (you can use BeautifulSoup or lxml)
        dasar_pelaksanaan = BeautifulSoup(spt_text.dasar_pelaksanaan, 'html.parser')
        items_dasar = [li.text for li in dasar_pelaksanaan.find_all('li')]
        tujuan_pelaksanaan = BeautifulSoup(spt_text.tujuan_pelaksanaan, 'html.parser')
        items_tujuan = [li.text for li in tujuan_pelaksanaan.find_all('li')]
        
        page_size = doc.sections[0]
        page_size.page_width = Mm(215.9)
        page_size.page_height = Mm(330.0)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.bottom_margin = Inches(1.0)
        #penambahan gambar kop
        doc.add_picture('static/img/KOP RS MANDALIKA 2024.png', width=Inches(7.0))
        add_content(doc=doc, content2=f'SURAT TUGAS', align='center', set_bold=True, set_underline=True)
        add_content(doc=doc, content2='Nomor : ${nomor_naskah}', align='center')
        # add_content(doc=doc, 'Lombok Tengah, ${tanggal_naskah}', align='right', space_after=6)
        table2 = doc.add_table(rows=4, cols=3)
        self.set_col_widths(table2)
        add_table_content(table=table2, row_index=0, col_index=0, content='Dasar')
        add_table_content(table=table2, row_index=0, col_index=1, content=':')
        for item in items_dasar:
            add_table_content(table=table2, row_index=0, col_index=2, content2=f'{item}', left_indent=5, style='List Number')
        
        add_table_content(table=table2, row_index=1, col_index=0, content='Kepada')
        add_table_content(table=table2, row_index=1, col_index=1, content=':')
        if pegawai:
            add_table_content(table=table2, row_index=1, col_index=2, content=f'Nama \t\t\t: {pegawai.full_name_2}')
            add_table_content(table=table2, row_index=1, col_index=2, content=f'Pangkat/Gol \t\t: {panggol_sdm.panggol if panggol_sdm is not None and hasattr(panggol_sdm, "panggol") else "-"}')
            add_table_content(table=table2, row_index=1, col_index=2, content=f'Jabatan/Profesi \t: {jabatan.nama_jabatan if jabatan is not None and hasattr(jabatan, "nama_jabatan") else "-"}')
            add_table_content(table=table2, row_index=1, col_index=2, content=f'NIP \t\t\t: {pegawai.profil_user.nip if hasattr(pegawai, "profil_user") else ""}')
        else:
            add_table_content(table=table2, row_index=1, col_index=2, content=f'Nama \t\t\t: -')
            add_table_content(table=table2, row_index=1, col_index=2, content=f'Pangkat/Gol \t\t: -')
            add_table_content(table=table2, row_index=1, col_index=2, content=f'Jabatan/Profesi \t: -')
            add_table_content(table=table2, row_index=1, col_index=2, content=f'NIP \t\t\t: -')
        table2.cell(2,0).merge(table2.cell(2,2))#merge row ketiga sampai kolom ketiga
        add_table_content(table=table2, row_index=2, col_index=0, content='MEMERINTAHKAN:', align='center')
        add_table_content(table=table2, row_index=3, col_index=0, content='Untuk')
        add_table_content(table=table2, row_index=3, col_index=1, content=':')
        for item in items_tujuan:
            add_table_content(table=table2, row_index=3, col_index=2, content2=f'{item}', left_indent=5, style='List Number 2')

        add_content(doc=doc, content='Lombok Tengah, ${tanggal_naskah}', space_before=30, left_indent=80)
        add_content(doc=doc, content='Direktur', left_indent=80, space_after=30)
        add_content(doc=doc, content='${ttd_pengirim}', left_indent=80, )
        add_content(doc=doc, content=f'{get_active_leader(penempatan.penempatan_level1).full_name_2}', left_indent=80, space_before=30)
        add_content(doc=doc, content=f'{panggol.panggol if hasattr(panggol, "panggol") else None}', left_indent=80)
        add_content(doc=doc, content=f'NIP. {get_active_leader(penempatan.penempatan_level1).profil_user.nip}', left_indent=80)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=pelimpahan_tugas-{pegawai.full_name}.docx'
        doc.save(response)

        return response


#SPT diklat untuk lebih dari satu orang
class LayananDiklatSPTDocxView2(View):
    def set_col_widths(self, table):
        widths = (Inches(2.5), Inches(0.3), Inches(9.5))
        # height = (Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                # row.cells[idx].height = heigh
                
    def set_col_table3_widths(self, table):
        widths = (Inches(2.5), Inches(0.3), Inches(0.5), Inches(9.5))
        # height = (Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.2))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                # row.cells[idx].height = heigh

    def get_object(self, id):
        try:
            data = RiwayatDiklat.objects.get(id=id)
            return data
        except Exception:
            return None

    def get(self, request, **kwargs):
        #Pengaturan ukuran kertas
        doc:Document=CreateDocument()
        id = kwargs.get('diklat_id')
        riwayatdiklat = self.get_object(id=id)
        spt_text = TextSPTDiklat.objects.filter(diklat=riwayatdiklat.usulan).first()
        jabatan = None
        panggol = None
        penempatan = None
        pegawai = riwayatdiklat.pegawai.all()
        detail_pegawai = riwayatdiklat.pegawai.first()
        if detail_pegawai:
            penempatan = RiwayatPenempatan.objects.filter(pegawai=detail_pegawai, status=True).last()
            panggol = get_active_leader(penempatan.penempatan_level1).riwayatpanggol_set.last() if hasattr(penempatan, 'penempatan_level1') and get_active_leader(penempatan.penempatan_level1) else None
        # Parse the HTML content (you can use BeautifulSoup or lxml)
        dasar_pelaksanaan = BeautifulSoup(spt_text.dasar_pelaksanaan, 'html.parser')
        items_dasar = [li.text for li in dasar_pelaksanaan.find_all('li')]
        tujuan_pelaksanaan = BeautifulSoup(spt_text.tujuan_pelaksanaan, 'html.parser')
        items_tujuan = [li.text for li in tujuan_pelaksanaan.find_all('li')]
        
        page_size = doc.sections[0]
        page_size.page_width = Mm(215.9)
        page_size.page_height = Mm(330.0)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.bottom_margin = Inches(1.0)
        #penambahan gambar kop
        doc.add_picture('static/img/KOP RS MANDALIKA 2024.png', width=Inches(7.0))
        add_content(doc=doc, content2=f'SURAT TUGAS', align='center', set_bold=True, set_underline=True)
        add_content(doc=doc, content2='Nomor : ${nomor_naskah}', align='center')
        # add_content(doc=doc, 'Lombok Tengah, ${tanggal_naskah}', align='right', space_after=6)
        table2 = doc.add_table(rows=4, cols=3)
        self.set_col_widths(table2)
        add_table_content(table=table2, row_index=0, col_index=0, content='Dasar')
        add_table_content(table=table2, row_index=0, col_index=1, content=':')
        for item in items_dasar:
            add_table_content(table=table2, row_index=0, col_index=2, content2=f'{item}', left_indent=5, style='List Number')
        
        add_table_content(table=table2, row_index=1, col_index=0, content='Kepada')
        add_table_content(table=table2, row_index=1, col_index=1, content=':')
        add_table_content(table=table2, row_index=1, col_index=2, content=f'(Daftar Terlampir)')
        
        table2.cell(2,0).merge(table2.cell(2,2))#merge row ketiga sampai kolom ketiga
        add_table_content(table=table2, row_index=2, col_index=0, content='MEMERINTAHKAN:', align='center')
        add_table_content(table=table2, row_index=3, col_index=0, content='Untuk')
        add_table_content(table=table2, row_index=3, col_index=1, content=':')
        for item in items_tujuan:
            add_table_content(table=table2, row_index=3, col_index=2, content2=f'{item}', left_indent=5, style='List Number 2')

        add_content(doc=doc, content='Lombok Tengah, ${tanggal_naskah}', space_before=20, left_indent=80)
        add_content(doc=doc, content='Direktur', left_indent=80, space_after=30)
        add_content(doc=doc, content='${ttd_pengirim}', left_indent=80, )
        add_content(doc=doc, content=f'{get_active_leader(penempatan.penempatan_level1).full_name_2}', left_indent=80, space_before=30)
        add_content(doc=doc, content=f'{panggol.panggol if hasattr(panggol, "panggol") else None}', left_indent=80)
        add_content(doc=doc, content=f'NIP. {get_active_leader(penempatan.penempatan_level1).profil_user.nip}', left_indent=80)

        doc.add_page_break()
        add_content(doc=doc, content='Lampiran:')
        add_content(doc=doc, content='Surat Tugas')
        add_content(doc=doc, content='Nomor: ${nomor_naskah}')
        rows = pegawai.count() + 1
        table3 = doc.add_table(rows=rows, cols=4)
        self.set_col_table3_widths(table3)
        add_table_content(table=table3, row_index=0, col_index=0, content='Kepada')
        add_table_content(table=table3, row_index=0, col_index=1, content=':')
        for i, p in enumerate(pegawai, start=0):
            add_table_content(table=table3, row_index=i, col_index=2, content=f'{i+1}.')
            add_table_content(table=table3, row_index=i, col_index=3, content=f'Nama \t\t\t: {p.full_name_2}')
            if p.riwayatjabatan_set.last():
                add_table_content(table=table3, row_index=i, col_index=3, content=f'Jabatan/Profesi \t: {p.riwayatjabatan_set.last().nama_jabatan}')
            else:
                add_table_content(table=table3, row_index=i, col_index=3, content=f'Jabatan/Profesi \t: -')
            if p.profil_user:
                add_table_content(table=table3, row_index=i, col_index=3, content=f'NIP \t\t\t: {p.profil_user.nip}')
            else:
                add_table_content(table=table3, row_index=i, col_index=3, content=f'NIP \t\t\t: -')
            if p.riwayatpanggol_set.last():
                add_table_content(table=table3, row_index=i, col_index=3, content=f'Pangkat/Gol \t\t: {p.riwayatpanggol_set.last().panggol}')
            else:
                add_table_content(table=table3, row_index=i, col_index=3, content=f'Pangkat/Gol \t\t: -')
            if p.riwayatjabatan_set.last():
                add_table_content(table=table3, row_index=i, col_index=3, content=f'Jabatan \t\t: {p.riwayatjabatan_set.last().detail_nama_jabatan}')
            else:
                add_table_content(table=table3, row_index=i, col_index=3, content=f'Jabatan \t\t: -')
        
        add_content(doc=doc, content='Lombok Tengah, ${tanggal_naskah}', space_before=20, left_indent=80)
        add_content(doc=doc, content='Direktur', left_indent=80, space_after=30)
        add_content(doc=doc, content='${ttd_pengirim}', left_indent=80, )
        add_content(doc=doc, content=f'{get_active_leader(penempatan.penempatan_level1).full_name_2}', left_indent=80, space_before=30)
        add_content(doc=doc, content=f'{panggol.panggol if hasattr(panggol, "panggol") else None}', left_indent=80)
        add_content(doc=doc, content=f'NIP. {get_active_leader(penempatan.penempatan_level1).profil_user.nip}', left_indent=80)
        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=spt-pelatihan-{detail_pegawai.full_name}.docx'
        doc.save(response)

        return response


class TextSPTDiklatView(View):
    def get_object(self, id):
        try:
            data=TextSPTDiklat.objects.get(id=id)
            return data
        except TextSPTDiklat.DoesNotExist:
            return None
        
    def get_diklat_object(self, id):
        try:
            data = RiwayatDiklat.objects.get(id=id)
            return data
        except RiwayatDiklat.DoesNotExist:
            return None

    def get(self, request, *args, **kwargs):
        id_diklat = kwargs.get('diklat_id')
        id = kwargs.get('id')
        get_form = request.GET.get('f')
        data_view = 'block'
        form_view = 'none'
        if get_form is not None:
            form_view = 'block'
            data_view = 'none'
        detail = self.get_object(id)
        riwayatdiklat = self.get_diklat_object(id_diklat)
        data=TextSPTDiklat.objects.filter(diklat=riwayatdiklat.usulan)
        form = TextSPTDiklatForm(instance=detail, diklat=riwayatdiklat.usulan)
        context={
            'id_layanan':id_diklat,
            'layanan':riwayatdiklat.usulan,
            'data':data,
            'form':form,
            'data_view':data_view,
            'form_view':form_view
        }
        return render(request, 'spt_text/spt_text_master.html', context)
    
    def post(self, request, **kwargs):
        id = kwargs.get('id')
        id_diklat = kwargs.get('diklat_id')
        detail = self.get_object(id)
        form = TextSPTDiklatForm(data=request.POST, instance=detail)
        if form.is_valid():
            form.save()
            messages.success(request, "Data berhasil disimpan!")
            return redirect(reverse('file_urls:text_spt_view', kwargs={'diklat_id':id_diklat}))
        messages.error(request, 'Datata gagal disimpan!')
        return redirect(reverse('file_urls:text_spt_view', kwargs={'diklat_id':id_diklat}))
    
# Menggunakan metode docxtpl

import os
from django.http import FileResponse, Http404
from django.contrib.auth.mixins import LoginRequiredMixin

from layanan.models import LayananSIP
from .services.sip_docx import (
    generate_permohonan_rekomendasi_sip,
    generate_surat_kecukupan_skp,
    generate_rekomendasi_sip,
)


class GeneratePermohonanRekomendasiSIPView(LoginRequiredMixin, View):
    def get(self, request, pk):
        if request.user.is_sip_admin:
            layanan_sip = (
                LayananSIP.objects
                .select_related(
                    "pegawai",
                    "pegawai__profil_user",
                    "layanan",
                    "ijazah",
                    "str_profesi",
                )
                .filter(pk=pk)
                .first()
            )
        else:
            layanan_sip = (
                LayananSIP.objects
                .select_related(
                    "pegawai",
                    "pegawai__profil_user",
                    "layanan",
                    "ijazah",
                    "str_profesi",
                )
                .filter(pk=pk, pegawai=request.user)
                .first()
            )

        if not layanan_sip:
            raise Http404("Data permohonan SIP tidak ditemukan.")

        file_path = generate_permohonan_rekomendasi_sip(layanan_sip)

        return FileResponse(
            open(file_path, "rb"),
            as_attachment=True,
            filename=os.path.basename(file_path),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


class GenerateSuratKecukupanSKPView(LoginRequiredMixin, View):
    def get(self, request, pk):
        if request.user.is_sip_admin:
            layanan_sip = (
                LayananSIP.objects
                .select_related(
                    "pegawai",
                    "pegawai__profil_user",
                    "layanan",
                    "ijazah",
                    "str_profesi",
                )
                .filter(pk=pk)
                .first()
            )
        else:
            layanan_sip = (
                LayananSIP.objects
            .select_related(
                "pegawai",
                "pegawai__profil_user",
                "layanan",
                "ijazah",
                "str_profesi",
            )
            .filter(pk=pk, pegawai=request.user)
            .first()
        )

        if not layanan_sip:
            raise Http404("Data permohonan SIP tidak ditemukan.")

        file_path = generate_surat_kecukupan_skp(layanan_sip)

        return FileResponse(
            open(file_path, "rb"),
            as_attachment=True,
            filename=os.path.basename(file_path),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        
        
class GenerateRekomendasiSIPView(LoginRequiredMixin, View):
    def get(self, request, pk):
        if request.user.is_sip_admin:
            layanan_sip = (
                LayananSIP.objects
                .select_related(
                    "pegawai",
                    "pegawai__profil_user",
                    "layanan",
                    "ijazah",
                    "str_profesi",
                )
                .filter(pk=pk)
                .first()
            )
        else:
            layanan_sip = (
                LayananSIP.objects
                .select_related(
                    "pegawai",
                    "pegawai__profil_user",
                    "layanan",
                    "ijazah",
                    "str_profesi",
                )
                .filter(pk=pk, pegawai=request.user)
                .first()
            )

        if not layanan_sip:
            raise Http404("Data permohonan SIP tidak ditemukan.")

        file_path = generate_rekomendasi_sip(layanan_sip)

        return FileResponse(
            open(file_path, "rb"),
            as_attachment=True,
            filename=os.path.basename(file_path),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
